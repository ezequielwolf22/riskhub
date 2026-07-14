"""Generacion de informes PDF y Excel (Risk Register, SoA, informes IA)."""
import io
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, Query, Request
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import (
    Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, PageBreak,
)
from sqlalchemy.orm import Session

from app.database import get_db
from app.i18n import get_lang, t as _t
from app.models import (
    AiConfig,
    Asset, Control, ControlImplementation, DPIA, DPIAStatus, Incident, IncidentStatus,
    Policy, PolicyStatus, ProcessingActivity, ReportBrandingConfig, Risk, RiskContext, RiskStatus,
    TreatmentTask, TaskStatus, User,
)
from app.security import filter_by_org, get_current_user
from app.services import report_ai_service

router = APIRouter(prefix="/api/reports", tags=["reports"])

# Colores brand por defecto (usados si no hay template configurado)
BRAND_PURPLE = colors.HexColor("#59008D")
BRAND_ORANGE = colors.HexColor("#D65200")
BRAND_GRAY1 = colors.HexColor("#262626")
BRAND_GRAY3 = colors.HexColor("#9D9D9D")
BRAND_GRAY5 = colors.HexColor("#E9E9E9")


@dataclass
class ReportBrand:
    """Configuracion de marca aplicada a la generacion del informe."""
    primary_color: str = "#59008D"
    secondary_color: str = "#D65200"
    font_family: str = "Helvetica"
    company_name: str = ""
    header_title: str = ""
    footer_text: str = ""
    cover_subtitle: str = ""
    logo_path: Optional[str] = None

    @property
    def primary(self):
        return colors.HexColor(self.primary_color)

    @property
    def secondary(self):
        return colors.HexColor(self.secondary_color)

    @property
    def font_bold(self):
        _bold_map = {
            "Helvetica": "Helvetica-Bold",
            "Times-Roman": "Times-Bold",
            "Courier": "Courier-Bold",
        }
        return _bold_map.get(self.font_family, "Helvetica-Bold")


def _load_brand(db: Session, org_id: int, report_type: str) -> ReportBrand:
    """Carga el template de marca para un tipo de informe.
    Prioridad: especifico del tipo > 'all' > defaults."""
    from pathlib import Path

    def _logo_root():
        p = Path("/srv/data/branding")
        if not p.exists():
            p = Path(__file__).parent.parent.parent / "data" / "branding"
        return p

    row = None
    if org_id:
        row = (
            db.query(ReportBrandingConfig)
            .filter(
                ReportBrandingConfig.organization_id == org_id,
                ReportBrandingConfig.report_type == report_type,
            )
            .first()
        )
        if not row:
            row = (
                db.query(ReportBrandingConfig)
                .filter(
                    ReportBrandingConfig.organization_id == org_id,
                    ReportBrandingConfig.report_type == "all",
                )
                .first()
            )

    if not row:
        return ReportBrand()

    logo_path = None
    if row.logo_filename:
        p = _logo_root() / row.logo_filename
        if p.exists():
            logo_path = str(p)

    return ReportBrand(
        primary_color=row.primary_color or "#59008D",
        secondary_color=row.secondary_color or "#D65200",
        font_family=row.font_family or "Helvetica",
        company_name=row.company_name or "",
        header_title=row.header_title or "",
        footer_text=row.footer_text or "",
        cover_subtitle=row.cover_subtitle or "",
        logo_path=logo_path,
    )


def _styles(brand: Optional[ReportBrand] = None):
    if brand is None:
        brand = ReportBrand()
    s = getSampleStyleSheet()
    s.add(ParagraphStyle("TitleBrand", parent=s["Title"], textColor=brand.primary,
                         fontName=brand.font_bold, fontSize=22, spaceAfter=6))
    s.add(ParagraphStyle("SubBrand", parent=s["Normal"], textColor=brand.secondary,
                         fontName=brand.font_bold, fontSize=12, spaceAfter=12))
    s.add(ParagraphStyle("H2Brand", parent=s["Heading2"], textColor=brand.primary,
                         fontName=brand.font_bold, fontSize=14, spaceBefore=12, spaceAfter=6))
    s.add(ParagraphStyle("BodyBrand", parent=s["Normal"], textColor=BRAND_GRAY1,
                         fontName=brand.font_family, fontSize=10, leading=13))
    return s


def _make_header_footer(brand: Optional[ReportBrand] = None, lang: str = "es"):
    """Devuelve una funcion de cabecera/pie personalizada con los colores y logo del template."""
    if brand is None:
        brand = ReportBrand()

    # Pre-calcular componentes RGB del gradiente
    try:
        r1 = int(brand.primary_color[1:3], 16) / 255
        g1 = int(brand.primary_color[3:5], 16) / 255
        b1 = int(brand.primary_color[5:7], 16) / 255
        r2 = int(brand.secondary_color[1:3], 16) / 255
        g2 = int(brand.secondary_color[3:5], 16) / 255
        b2 = int(brand.secondary_color[5:7], 16) / 255
    except (ValueError, IndexError):
        r1, g1, b1 = 0x59 / 255, 0.0, 0x8D / 255
        r2, g2, b2 = 0xD6 / 255, 0x52 / 255, 0.0

    footer_label = brand.footer_text or (brand.company_name + " - RiskHub" if brand.company_name else "RiskHub")
    logo_path = brand.logo_path

    def _header_footer_fn(canvas, doc):
        canvas.saveState()
        w, h = A4
        bar_y = h - 8 * mm
        steps = 60
        for i in range(steps):
            t = i / (steps - 1)
            canvas.setFillColorRGB(r1 + t * (r2 - r1), g1 + t * (g2 - g1), b1 + t * (b2 - b1))
            canvas.rect(i * w / steps, bar_y, w / steps + 0.5, 4 * mm, stroke=0, fill=1)

        # Logo en cabecera (esquina superior derecha)
        if logo_path:
            try:
                from reportlab.lib.utils import ImageReader
                img = ImageReader(logo_path)
                img_w, img_h = img.getSize()
                max_h = 6 * mm
                scale = max_h / img_h
                canvas.drawImage(
                    logo_path, w - 20 * mm - img_w * scale, bar_y - 0.5 * mm,
                    width=img_w * scale, height=img_h * scale,
                    mask="auto", preserveAspectRatio=True,
                )
            except Exception:
                pass

        canvas.setFont(brand.font_family, 8)
        canvas.setFillColor(BRAND_GRAY3)
        canvas.drawString(20 * mm, 10 * mm,
                          f"{footer_label} - {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        canvas.drawRightString(w - 20 * mm, 10 * mm, _t("reports.page_label", lang, page=doc.page))
        canvas.restoreState()

    return _header_footer_fn


# Funcion legacy sin brand — para compatibilidad interna
def _header_footer(canvas, doc):
    _make_header_footer()(canvas, doc)


def _pdf_response(elements, filename: str, brand: Optional[ReportBrand] = None, lang: str = "es"):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=20*mm, rightMargin=20*mm,
                            topMargin=20*mm, bottomMargin=20*mm)
    hf = _make_header_footer(brand, lang)
    doc.build(elements, onFirstPage=hf, onLaterPages=hf)
    buf.seek(0)
    return StreamingResponse(
        buf, media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@router.get("/risk-register")
def risk_register(request: Request, db: Session = Depends(get_db),
                  current_user: User = Depends(get_current_user)):
    """Risk Register (ISO 27005)."""
    lang = get_lang(request)
    brand = _load_brand(db, current_user.organization_id, "risk_register")
    BRAND_PURPLE = brand.primary
    BRAND_ORANGE = brand.secondary
    s = _styles(brand)
    el = []
    el.append(Paragraph(_t("reports.risk_register.title", lang), s["TitleBrand"]))
    el.append(Paragraph(_t("reports.risk_register.subtitle", lang), s["SubBrand"]))
    el.append(Spacer(1, 6))

    ctx = filter_by_org(db.query(RiskContext), RiskContext, current_user).first()
    if ctx:
        el.append(Paragraph(
            f"<b>{_t('reports.risk_register.organization_label', lang)}</b> {ctx.organization_name or '-'}",
            s["BodyBrand"]))
        el.append(Paragraph(
            f"<b>{_t('reports.risk_register.scope_label', lang)}</b> {ctx.scope or '-'}",
            s["BodyBrand"]))
        el.append(Paragraph(
            f"<b>{_t('reports.risk_register.appetite_label', lang)}</b> "
            f"{_t('reports.risk_register.appetite_level', lang, level=ctx.risk_appetite)}",
            s["BodyBrand"]))
        el.append(Spacer(1, 12))

    risks = filter_by_org(db.query(Risk), Risk, current_user).order_by(Risk.residual_level.desc()).all()
    data = [[
        _t("reports.risk_register.table_header_code", lang),
        _t("reports.risk_register.table_header_asset", lang),
        _t("reports.risk_register.table_header_threat", lang),
        _t("reports.risk_register.table_header_inherent", lang),
        _t("reports.risk_register.table_header_residual", lang),
        _t("reports.risk_register.table_header_status", lang),
        _t("reports.risk_register.table_header_treatment", lang),
        _t("reports.risk_register.table_header_owner", lang),
    ]]
    for r in risks:
        owner_name = r.owner.full_name if r.owner else "-"
        data.append([
            r.code,
            r.asset.name if r.asset else "-",
            r.threat.name if r.threat else "-",
            str(r.inherent_level),
            str(r.residual_level),
            r.status.value if r.status else "-",
            r.treatment_option.value if r.treatment_option else "-",
            owner_name,
        ])
    if len(data) > 1:
        t = Table(data, repeatRows=1, colWidths=[18*mm, 36*mm, 44*mm, 10*mm, 10*mm, 18*mm, 20*mm, 24*mm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), BRAND_PURPLE),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, BRAND_GRAY5]),
            ("GRID", (0, 0), (-1, -1), 0.25, BRAND_GRAY3),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        el.append(t)
    else:
        el.append(Paragraph(_t("reports.risk_register.no_risks", lang), s["BodyBrand"]))

    # Pagina de metodologia
    el.append(PageBreak())
    el.append(Paragraph(_t("reports.risk_register.methodology_title", lang), s["H2Brand"]))
    el.append(Spacer(1, 4))
    el.append(Paragraph(_t("reports.risk_register.methodology_body", lang), s["BodyBrand"]))
    el.append(Spacer(1, 8))
    el.append(Paragraph(_t("reports.risk_register.scale_title", lang), s["H2Brand"]))
    level_data = [
        [
            _t("reports.risk_register.scale_header_level", lang),
            _t("reports.risk_register.scale_header_range", lang),
            _t("reports.risk_register.scale_header_description", lang),
            _t("reports.risk_register.scale_header_action", lang),
        ],
        [_t("reports.risk_register.level_critical", lang), "7-8",
         _t("reports.risk_register.level_critical_desc", lang),
         _t("reports.risk_register.level_critical_action", lang)],
        [_t("reports.risk_register.level_high", lang), "5-6",
         _t("reports.risk_register.level_high_desc", lang),
         _t("reports.risk_register.level_high_action", lang)],
        [_t("reports.risk_register.level_medium", lang), "3-4",
         _t("reports.risk_register.level_medium_desc", lang),
         _t("reports.risk_register.level_medium_action", lang)],
        [_t("reports.risk_register.level_low", lang), "1-2",
         _t("reports.risk_register.level_low_desc", lang),
         _t("reports.risk_register.level_low_action", lang)],
        [_t("reports.risk_register.level_negligible", lang), "0",
         _t("reports.risk_register.level_negligible_desc", lang),
         _t("reports.risk_register.level_negligible_action", lang)],
    ]
    level_colors = [BRAND_GRAY5, colors.HexColor("#FCA5A5"), colors.HexColor("#FED7AA"),
                    colors.HexColor("#FEF9C3"), colors.HexColor("#DCFCE7"), colors.white]
    t2 = Table(level_data, repeatRows=1, colWidths=[25*mm, 18*mm, 80*mm, 47*mm])
    t2_style = [
        ("BACKGROUND", (0, 0), (-1, 0), BRAND_PURPLE),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.25, BRAND_GRAY3),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]
    for idx, bg in enumerate(level_colors[1:], 1):
        t2_style.append(("BACKGROUND", (0, idx), (-1, idx), bg))
    t2.setStyle(TableStyle(t2_style))
    el.append(t2)
    el.append(Spacer(1, 8))
    el.append(Paragraph(_t("reports.risk_register.treatment_title", lang), s["H2Brand"]))
    treat_data = [
        [
            _t("reports.risk_register.treatment_header_option", lang),
            _t("reports.risk_register.treatment_header_description", lang),
        ],
        [_t("reports.risk_register.treatment_modification", lang),
         _t("reports.risk_register.treatment_modification_desc", lang)],
        [_t("reports.risk_register.treatment_retention", lang),
         _t("reports.risk_register.treatment_retention_desc", lang)],
        [_t("reports.risk_register.treatment_transfer", lang),
         _t("reports.risk_register.treatment_transfer_desc", lang)],
        [_t("reports.risk_register.treatment_avoidance", lang),
         _t("reports.risk_register.treatment_avoidance_desc", lang)],
        [_t("reports.risk_register.treatment_sharing", lang),
         _t("reports.risk_register.treatment_sharing_desc", lang)],
    ]
    t3 = Table(treat_data, repeatRows=1, colWidths=[35*mm, 135*mm])
    t3.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), BRAND_ORANGE),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, BRAND_GRAY5]),
        ("GRID", (0, 0), (-1, -1), 0.25, BRAND_GRAY3),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    el.append(t3)

    return _pdf_response(el, "risk_register.pdf", brand, lang)


@router.get("/soa")
def statement_of_applicability(request: Request, db: Session = Depends(get_db),
                               current_user: User = Depends(get_current_user)):
    """Statement of Applicability ISO 27001/27002 — completo con todos los campos normativos."""
    lang = get_lang(request)
    brand = _load_brand(db, current_user.organization_id, "soa")
    BRAND_PURPLE = brand.primary
    BRAND_ORANGE = brand.secondary
    s = _styles(brand)
    el = []
    now_str = datetime.now().strftime("%d/%m/%Y")
    ctx = filter_by_org(db.query(RiskContext), RiskContext, current_user).first()
    org_name = (ctx.organization_name if ctx else None) or _t("reports.soa.default_org_name", lang)
    scope = (ctx.scope if ctx else None) or _t("reports.soa.default_scope", lang)
    appetite = ctx.risk_appetite if ctx else 3

    # ── PORTADA ──────────────────────────────────────────────────────────────
    el.append(Spacer(1, 30*mm))
    el.append(Paragraph(_t("reports.soa.title", lang), s["TitleBrand"]))
    el.append(Paragraph(_t("reports.soa.subtitle", lang), s["SubBrand"]))
    el.append(Spacer(1, 8))
    el.append(Paragraph(f"<b>{_t('reports.soa.organization_label', lang)}</b> {_safe(org_name)}", s["BodyBrand"]))
    el.append(Paragraph(f"<b>{_t('reports.soa.standard_label', lang)}</b> {_t('reports.soa.standard_value', lang)}", s["BodyBrand"]))
    el.append(Paragraph(f"<b>{_t('reports.soa.issue_date_label', lang)}</b> {now_str}", s["BodyBrand"]))
    el.append(Paragraph(f"<b>{_t('reports.soa.version_label', lang)}</b> {_t('reports.soa.version_value', lang)}", s["BodyBrand"]))
    el.append(Spacer(1, 16))

    # Alcance y contexto
    el.append(Paragraph(_t("reports.soa.scope_title", lang), s["H2Brand"]))
    el.append(Paragraph(_safe(scope), s["BodyBrand"]))
    el.append(Paragraph(f"<b>{_t('reports.soa.appetite_label', lang)}</b> {_t('reports.soa.appetite_level', lang, level=appetite)}", s["BodyBrand"]))
    el.append(PageBreak())

    # ── TABLA DE CONTROL DE VERSIONES ────────────────────────────────────────
    el.append(Paragraph(_t("reports.soa.version_control_title", lang), s["H2Brand"]))
    ver_data = [
        [_t("reports.soa.version_table_header_version", lang), _t("reports.soa.version_table_header_date", lang),
         _t("reports.soa.version_table_header_author", lang), _t("reports.soa.version_table_header_description", lang)],
        ["1.0", now_str, _t("reports.soa.version_table_initial_author", lang), _t("reports.soa.version_table_initial_description", lang)],
    ]
    ver_t = Table(ver_data, colWidths=[20*mm, 30*mm, 60*mm, 60*mm])
    ver_t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), BRAND_PURPLE),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.3, BRAND_GRAY3),
    ]))
    el.append(ver_t)
    el.append(Spacer(1, 12))

    # ── ESTADISTICAS ─────────────────────────────────────────────────────────
    impls = filter_by_org(db.query(ControlImplementation), ControlImplementation, current_user).all()
    by_control = {}
    for imp in impls:
        by_control.setdefault(imp.control_id, []).append(imp)

    controls = db.query(Control).order_by(Control.code).all()
    total_c = len(controls)
    total_impl = sum(1 for c in controls if by_control.get(c.id))
    implemented_count = sum(1 for imp in impls if imp.status and imp.status.value == "implemented")
    partial_count = sum(1 for imp in impls if imp.status and imp.status.value == "partial")
    excluded_count = sum(1 for imp in impls if imp.exclusion_justification)

    el.append(Paragraph(_t("reports.soa.executive_summary_title", lang), s["H2Brand"]))
    stats_data = [
        [_t("reports.soa.stats_header_indicator", lang), _t("reports.soa.stats_header_value", lang)],
        [_t("reports.soa.stats_total_controls", lang), str(total_c)],
        [_t("reports.soa.stats_active_implementation", lang), str(total_impl)],
        [_t("reports.soa.stats_implemented", lang), str(implemented_count)],
        [_t("reports.soa.stats_partial", lang), str(partial_count)],
        [_t("reports.soa.stats_excluded", lang), str(excluded_count)],
        [_t("reports.soa.stats_not_implemented", lang), str(total_c - total_impl)],
    ]
    stats_t = Table(stats_data, colWidths=[120*mm, 50*mm])
    stats_t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), BRAND_PURPLE),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, BRAND_GRAY5]),
        ("GRID", (0, 0), (-1, -1), 0.3, BRAND_GRAY3),
        ("ALIGN", (1, 0), (1, -1), "CENTER"),
    ]))
    el.append(stats_t)
    el.append(PageBreak())

    # ── TABLA COMPLETA DE CONTROLES ───────────────────────────────────────────
    el.append(Paragraph(_t("reports.soa.detail_title", lang), s["H2Brand"]))
    el.append(Paragraph(_t("reports.soa.detail_description", lang), s["BodyBrand"]))
    el.append(Spacer(1, 6))

    _STATUS_LABELS = {
        "implemented": _t("reports.soa.status_implemented", lang),
        "partial": _t("reports.soa.status_partial", lang),
        "planned": _t("reports.soa.status_planned", lang),
        "not_implemented": _t("reports.soa.status_not_implemented", lang),
    }
    _REASON_LABELS = {
        "legal": _t("reports.soa.reason_legal", lang),
        "contractual": _t("reports.soa.reason_contractual", lang),
        "risk": _t("reports.soa.reason_risk", lang),
        "best_practice": _t("reports.soa.reason_best_practice", lang),
    }

    data = [[
        _t("reports.soa.table_header_ctrl", lang), _t("reports.soa.table_header_control_name", lang),
        _t("reports.soa.table_header_applicable", lang), _t("reports.soa.table_header_status", lang),
        _t("reports.soa.table_header_maturity", lang), _t("reports.soa.table_header_inclusion_reason", lang),
        _t("reports.soa.table_header_exclusion_reason", lang), _t("reports.soa.table_header_evidence", lang),
        _t("reports.soa.table_header_last_soa_review", lang), _t("reports.soa.table_header_next_review", lang),
    ]]

    _YES = _t("reports.soa.yes", lang)
    _NO = _t("reports.soa.no", lang)

    for c in controls:
        ims = by_control.get(c.id, [])
        if ims:
            for imp in ims:
                refs = imp.evidence_refs if isinstance(imp.evidence_refs, list) else []
                ev_str = "\n".join(
                    (r.get("title") or "")[:28] for r in refs[:3] if r.get("title")
                ) or (imp.evidence or "-")[:40]
                soa_rev = imp.soa_reviewed_at.strftime("%d/%m/%y") if imp.soa_reviewed_at else "-"
                next_rev = imp.next_review.strftime("%d/%m/%y") if imp.next_review else "-"
                data.append([
                    c.code,
                    _safe(c.name, 42),
                    _YES,
                    _STATUS_LABELS.get(imp.status.value if imp.status else "", "-"),
                    f"{imp.maturity or 0}/5",
                    _REASON_LABELS.get(imp.inclusion_reason or "", imp.inclusion_reason or "-"),
                    _safe(imp.exclusion_justification, 35) if imp.exclusion_justification else "-",
                    _safe(ev_str, 40),
                    soa_rev,
                    next_rev,
                ])
        else:
            data.append([c.code, _safe(c.name, 42), _NO, "-", "-", "-", "-", "-", "-", "-"])

    # Columnas ajustadas para A4 landscape-ish en puntos
    col_w = [13*mm, 48*mm, 10*mm, 20*mm, 10*mm, 24*mm, 24*mm, 30*mm, 18*mm, 18*mm]
    t = Table(data, repeatRows=1, colWidths=col_w)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), BRAND_PURPLE),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 6),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, BRAND_GRAY5]),
        ("GRID", (0, 0), (-1, -1), 0.2, BRAND_GRAY3),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("WORDWRAP", (1, 1), (1, -1), True),
    ]))
    el.append(t)

    # ── SECCION DE FIRMA Y APROBACION ─────────────────────────────────────────
    el.append(PageBreak())
    el.append(Paragraph(_t("reports.soa.approval_title", lang), s["H2Brand"]))
    el.append(Paragraph(_t("reports.soa.approval_description", lang), s["BodyBrand"]))
    el.append(Spacer(1, 20))
    firma_data = [
        [_t("reports.soa.signature_header_role", lang), _t("reports.soa.signature_header_name", lang),
         _t("reports.soa.signature_header_date", lang), _t("reports.soa.signature_header_signature", lang)],
        [_t("reports.soa.signature_role_security_officer", lang), "_" * 30, "_" * 15, "_" * 20],
        [_t("reports.soa.signature_role_ciso", lang), "_" * 30, "_" * 15, "_" * 20],
        [_t("reports.soa.signature_role_internal_auditor", lang), "_" * 30, "_" * 15, "_" * 20],
    ]
    firma_t = Table(firma_data, colWidths=[45*mm, 55*mm, 30*mm, 40*mm])
    firma_t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), BRAND_PURPLE),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, BRAND_GRAY5]),
        ("GRID", (0, 0), (-1, -1), 0.3, BRAND_GRAY3),
        ("ROWHEIGHT", (0, 1), (-1, -1), 18*mm),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    el.append(firma_t)
    el.append(Spacer(1, 10))
    el.append(Paragraph(_t("reports.soa.footer_note", lang), s["BodyBrand"]))

    return _pdf_response(el, f"SOA_{org_name.replace(' ','_')}_{datetime.now().strftime('%Y%m%d')}.pdf", brand, lang)


# ============================================================
# Excel — Risk Register completo
# ============================================================

@router.get("/risk-register-excel")
def risk_register_excel(request: Request, db: Session = Depends(get_db),
                        current_user: User = Depends(get_current_user)):
    """Exporta el Risk Register completo a Excel (.xlsx)."""
    lang = get_lang(request)
    import openpyxl
    from openpyxl.styles import (
        Alignment, Border, Font, PatternFill, Side,
    )
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()

    # ---- Hoja 1: Riesgos ----
    ws = wb.active
    ws.title = _t("reports.risk_register_excel.sheet_risk_register", lang)

    hdr_fill = PatternFill("solid", fgColor="59008D")
    hdr_font = Font(color="FFFFFF", bold=True, size=10)
    alt_fill = PatternFill("solid", fgColor="F3E8FF")
    thin = Side(style="thin", color="D1D5DB")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    headers = [
        _t("reports.risk_register_excel.rr_header_code", lang), _t("reports.risk_register_excel.rr_header_asset", lang),
        _t("reports.risk_register_excel.rr_header_asset_type", lang), _t("reports.risk_register_excel.rr_header_threat", lang),
        _t("reports.risk_register_excel.rr_header_threat_category", lang),
        _t("reports.risk_register_excel.rr_header_inherent_likelihood", lang), _t("reports.risk_register_excel.rr_header_inherent_consequence", lang),
        _t("reports.risk_register_excel.rr_header_inherent_level", lang),
        _t("reports.risk_register_excel.rr_header_residual_likelihood", lang), _t("reports.risk_register_excel.rr_header_residual_consequence", lang),
        _t("reports.risk_register_excel.rr_header_residual_level", lang),
        _t("reports.risk_register_excel.rr_header_status", lang), _t("reports.risk_register_excel.rr_header_treatment", lang), _t("reports.risk_register_excel.rr_header_due_date", lang),
        _t("reports.risk_register_excel.rr_header_description", lang), _t("reports.risk_register_excel.rr_header_treatment_plan", lang),
    ]
    ws.append(headers)
    for col_idx, _ in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border

    risks = filter_by_org(db.query(Risk), Risk, current_user).order_by(Risk.residual_level.desc()).all()
    for i, r in enumerate(risks, 2):
        row = [
            r.code,
            r.asset.name if r.asset else "-",
            r.asset.asset_type.value if r.asset and r.asset.asset_type else "-",
            r.threat.name if r.threat else "-",
            r.threat.category if r.threat else "-",
            r.inherent_likelihood,
            r.inherent_consequence,
            r.inherent_level,
            r.residual_likelihood,
            r.residual_consequence,
            r.residual_level,
            r.status.value if r.status else "-",
            r.treatment_option.value if r.treatment_option else "-",
            r.treatment_due_date.strftime("%Y-%m-%d") if r.treatment_due_date else "-",
            r.description or "",
            r.treatment_plan or "",
        ]
        ws.append(row)
        row_fill = alt_fill if i % 2 == 0 else None
        for col_idx in range(1, len(headers) + 1):
            cell = ws.cell(row=i, column=col_idx)
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            if row_fill:
                cell.fill = row_fill
            # Colorear nivel residual
            if col_idx == 11 and isinstance(cell.value, int):
                if cell.value >= 7:
                    cell.fill = PatternFill("solid", fgColor="FCA5A5")
                elif cell.value >= 5:
                    cell.fill = PatternFill("solid", fgColor="FED7AA")
                elif cell.value >= 3:
                    cell.fill = PatternFill("solid", fgColor="FEF9C3")

    col_widths = [12, 24, 18, 30, 18, 9, 9, 9, 9, 9, 9, 14, 16, 14, 40, 40]
    for idx, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(idx)].width = w
    ws.row_dimensions[1].height = 32
    ws.freeze_panes = "A2"

    # ---- Hoja 2: Activos ----
    ws2 = wb.create_sheet(_t("reports.risk_register_excel.sheet_assets", lang))
    ws2.append([
        _t("reports.risk_register_excel.assets_header_code", lang), _t("reports.risk_register_excel.assets_header_name", lang),
        _t("reports.risk_register_excel.assets_header_type", lang), _t("reports.risk_register_excel.assets_header_classification", lang),
        _t("reports.risk_register_excel.assets_header_location", lang), _t("reports.risk_register_excel.assets_header_c", lang),
        _t("reports.risk_register_excel.assets_header_i", lang), _t("reports.risk_register_excel.assets_header_a", lang),
        _t("reports.risk_register_excel.assets_header_max_value", lang),
    ])
    for col_idx in range(1, 10):
        c = ws2.cell(row=1, column=col_idx)
        c.fill = hdr_fill
        c.font = hdr_font
        c.border = border
        c.alignment = Alignment(horizontal="center")

    assets = filter_by_org(db.query(Asset), Asset, current_user).order_by(Asset.code).all()
    for i, a in enumerate(assets, 2):
        ws2.append([
            a.code, a.name,
            a.asset_type.value if a.asset_type else "-",
            a.classification or "-",
            a.location or "-",
            a.value_confidentiality, a.value_integrity, a.value_availability,
            a.value_max,
        ])
        for col_idx in range(1, 10):
            ws2.cell(row=i, column=col_idx).border = border

    for idx, w in enumerate([12, 28, 20, 16, 16, 6, 6, 6, 10], 1):
        ws2.column_dimensions[get_column_letter(idx)].width = w

    # ---- Hoja 3: Controles ----
    ws3 = wb.create_sheet(_t("reports.risk_register_excel.sheet_controls", lang))
    ws3.append([
        _t("reports.risk_register_excel.controls_header_control", lang), _t("reports.risk_register_excel.controls_header_name", lang),
        _t("reports.risk_register_excel.controls_header_theme", lang), _t("reports.risk_register_excel.controls_header_status", lang),
        _t("reports.risk_register_excel.controls_header_maturity", lang), _t("reports.risk_register_excel.controls_header_description", lang),
    ])
    for col_idx in range(1, 7):
        c = ws3.cell(row=1, column=col_idx)
        c.fill = hdr_fill
        c.font = hdr_font
        c.border = border
        c.alignment = Alignment(horizontal="center")

    impls = filter_by_org(db.query(ControlImplementation), ControlImplementation, current_user).all()
    for i, imp in enumerate(impls, 2):
        ws3.append([
            imp.control.code if imp.control else "-",
            imp.name,
            imp.control.theme if imp.control else "-",
            imp.status.value if imp.status else "-",
            imp.maturity,
            imp.description or "",
        ])
        for col_idx in range(1, 7):
            ws3.cell(row=i, column=col_idx).border = border

    for idx, w in enumerate([10, 32, 18, 18, 10, 40], 1):
        ws3.column_dimensions[get_column_letter(idx)].width = w

    # ---- Hoja 4: Resumen ----
    ws4 = wb.create_sheet(_t("reports.risk_register_excel.sheet_summary", lang))
    ctx = filter_by_org(db.query(RiskContext), RiskContext, current_user).first()
    stat_fill = PatternFill("solid", fgColor="EDE9FE")
    rows_summary = [
        (_t("reports.risk_register_excel.summary_organization", lang), ctx.organization_name if ctx else "-"),
        (_t("reports.risk_register_excel.summary_scope", lang), (ctx.scope or "-") if ctx else "-"),
        (_t("reports.risk_register_excel.summary_appetite", lang), _t("reports.risk_register_excel.summary_appetite_level", lang, level=ctx.risk_appetite if ctx else "-")),
        (_t("reports.risk_register_excel.summary_export_date", lang), datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("", ""),
        (_t("reports.risk_register_excel.summary_total_risks", lang), len(risks)),
        (_t("reports.risk_register_excel.summary_critical_level", lang), sum(1 for r in risks if r.residual_level >= 7)),
        (_t("reports.risk_register_excel.summary_high_level", lang), sum(1 for r in risks if 5 <= r.residual_level < 7)),
        (_t("reports.risk_register_excel.summary_medium_level", lang), sum(1 for r in risks if 3 <= r.residual_level < 5)),
        (_t("reports.risk_register_excel.summary_low_level", lang), sum(1 for r in risks if r.residual_level < 3)),
        (_t("reports.risk_register_excel.summary_with_treatment_plan", lang), sum(1 for r in risks if r.treatment_option)),
        (_t("reports.risk_register_excel.summary_accepted", lang), sum(1 for r in risks if r.status == RiskStatus.ACCEPTED)),
        ("", ""),
        (_t("reports.risk_register_excel.summary_total_assets", lang), len(assets)),
        (_t("reports.risk_register_excel.summary_total_controls_implemented", lang), len(impls)),
    ]
    for label, value in rows_summary:
        ws4.append([label, value])
    for row_idx in range(1, len(rows_summary) + 1):
        for col_idx in [1, 2]:
            cell = ws4.cell(row=row_idx, column=col_idx)
            cell.border = border
            if col_idx == 1 and ws4.cell(row=row_idx, column=1).value:
                cell.font = Font(bold=True)
            if row_idx in [1, 2, 3, 4]:
                cell.fill = stat_fill

    ws4.column_dimensions["A"].width = 32
    ws4.column_dimensions["B"].width = 30

    # ---- Hoja 5: Dashboard visual ─────────────────────────────────────────
    from openpyxl.chart import BarChart, PieChart, Reference, Series as ChartSeries  # noqa
    ws5 = wb.create_sheet("Dashboard", 0)   # primera hoja
    wb.active = ws5

    # Titulo
    ws5["A1"] = _t("reports.risk_register_excel.dashboard_title", lang)
    ws5["A1"].font = Font(size=16, bold=True, color="59008D")
    ws5["A2"] = f"{ctx.organization_name if ctx else _t('reports.risk_register_excel.dashboard_org_fallback', lang)} — {datetime.now().strftime('%d/%m/%Y')}"
    ws5["A2"].font = Font(size=11, color="9D9D9D")

    # KPIs en fila
    kpi_col = 1
    for label, val, bg in [
        (_t("reports.risk_register_excel.kpi_total_risks", lang), len(risks), "EDE9FE"),
        (_t("reports.risk_register_excel.kpi_high_risks", lang), sum(1 for r in risks if r.residual_level >= 5), "FEE2E2"),
        (_t("reports.risk_register_excel.kpi_controls_implemented", lang), sum(1 for imp in impls if imp.status and imp.status.value == "implemented"), "D1FAE5"),
        (_t("reports.risk_register_excel.kpi_average_maturity", lang), f"{round(sum(imp.maturity or 0 for imp in impls)/max(1, len(impls)), 1)}/5", "FEF9C3"),
    ]:
        ws5.cell(row=4, column=kpi_col, value=label).font = Font(bold=True, size=9, color="59008D")
        cell_val = ws5.cell(row=5, column=kpi_col, value=val)
        cell_val.font = Font(size=18, bold=True)
        cell_val.fill = PatternFill("solid", fgColor=bg)
        cell_val.alignment = Alignment(horizontal="center")
        ws5.column_dimensions[get_column_letter(kpi_col)].width = 22
        kpi_col += 2

    # Datos para graficos
    chart_row = 8
    ws5.cell(row=chart_row, column=1, value=_t("reports.risk_register_excel.chart_residual_level", lang))
    ws5.cell(row=chart_row, column=2, value=_t("reports.risk_register_excel.chart_quantity", lang))
    for lbl, fn in [(_t("reports.risk_register_excel.chart_level_critical", lang), lambda r: r.residual_level >= 7),
                    (_t("reports.risk_register_excel.chart_level_high", lang), lambda r: 5 <= r.residual_level < 7),
                    (_t("reports.risk_register_excel.chart_level_medium", lang), lambda r: 3 <= r.residual_level < 5),
                    (_t("reports.risk_register_excel.chart_level_low", lang), lambda r: r.residual_level < 3)]:
        chart_row += 1
        ws5.cell(row=chart_row, column=1, value=lbl)
        ws5.cell(row=chart_row, column=2, value=sum(1 for r in risks if fn(r)))

    try:
        bar = BarChart()
        bar.title = _t("reports.risk_register_excel.bar_chart_title", lang)
        bar.style = 10
        bar.y_axis.title = _t("reports.risk_register_excel.bar_chart_y_axis", lang)
        bar.x_axis.title = _t("reports.risk_register_excel.bar_chart_x_axis", lang)
        bar.shape = 4
        data_ref = Reference(ws5, min_col=2, min_row=8, max_row=12)
        cats = Reference(ws5, min_col=1, min_row=9, max_row=12)
        bar.add_data(data_ref, titles_from_data=True)
        bar.set_categories(cats)
        bar.width = 15; bar.height = 10
        ws5.add_chart(bar, "E4")

        pie = PieChart()
        pie.title = _t("reports.risk_register_excel.pie_chart_title", lang)
        pie.style = 10
        st_row = chart_row + 2
        ws5.cell(row=st_row, column=1, value=_t("reports.risk_register_excel.pie_chart_status_label", lang))
        ws5.cell(row=st_row, column=2, value=_t("reports.risk_register_excel.pie_chart_n_label", lang))
        for st_lbl, st_val in [
            (_t("reports.risk_register_excel.status_implemented", lang), sum(1 for i in impls if i.status and i.status.value == "implemented")),
            (_t("reports.risk_register_excel.status_partial", lang), sum(1 for i in impls if i.status and i.status.value == "partial")),
            (_t("reports.risk_register_excel.status_planned", lang), sum(1 for i in impls if i.status and i.status.value == "planned")),
            (_t("reports.risk_register_excel.status_not_implemented", lang), sum(1 for i in impls if i.status and i.status.value == "not_implemented")),
        ]:
            st_row += 1
            ws5.cell(row=st_row, column=1, value=st_lbl)
            ws5.cell(row=st_row, column=2, value=st_val)
        pie_data = Reference(ws5, min_col=2, min_row=chart_row + 2, max_row=st_row)
        pie_cats = Reference(ws5, min_col=1, min_row=chart_row + 3, max_row=st_row)
        pie.add_data(pie_data, titles_from_data=True)
        pie.set_categories(pie_cats)
        pie.width = 15; pie.height = 10
        ws5.add_chart(pie, "M4")
    except Exception:
        pass  # charts opcionales — no bloquear si hay error

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    fname = f"riskhub_dashboard_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={fname}"},
    )


# ============================================================
# Informe de Revision por la Direccion — ISO 27001:2022 cl. 9.3 + ENS
# ============================================================

@router.get("/management-review")
def management_review(
    request: Request,
    format: str = Query("pdf", regex="^(pdf|excel|word)$"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Informe de Revision por la Direccion (ISO 27001:2022 cl. 9.3 / ENS).

    Soporta format=pdf | excel | word.
    """
    lang = get_lang(request)
    from app.models import AuditProgram, NonConformity, Supplier  # noqa
    ctx = filter_by_org(db.query(RiskContext), RiskContext, current_user).first()
    org_name = (ctx.organization_name if ctx else None) or "Organizacion"
    scope = (ctx.scope if ctx else None) or ""
    boundaries = (ctx.boundaries if ctx else None) or ""
    now = datetime.now(timezone.utc)
    now_str = now.strftime("%d/%m/%Y")

    # Recopilar datos reales
    risks = filter_by_org(db.query(Risk), Risk, current_user).all()
    impls = filter_by_org(db.query(ControlImplementation), ControlImplementation, current_user).all()
    incidents_all = filter_by_org(db.query(Incident), Incident, current_user).all()
    tasks_all = filter_by_org(db.query(TreatmentTask), TreatmentTask, current_user).all()
    policies = filter_by_org(db.query(Policy), Policy, current_user).all()
    suppliers = filter_by_org(db.query(Supplier), Supplier, current_user).all()

    total_risks = len(risks)
    high_risks = sum(1 for r in risks if r.residual_level >= 5)
    untreated = sum(1 for r in risks if not r.treatment_option and r.residual_level >= 5)
    impl_pct = round(sum(1 for i in impls if i.status and i.status.value == "implemented") / max(1, len(impls)) * 100)
    avg_mat = round(sum(i.maturity or 0 for i in impls) / max(1, len(impls)), 1)
    open_inc = sum(1 for i in incidents_all if i.status and i.status.value in ("open", "in_progress"))
    p1p2 = sum(1 for i in incidents_all if i.severity and i.severity.value in ("p1", "p2"))
    tasks_overdue = sum(
        1 for t in tasks_all
        if t.due_date and t.due_date.replace(tzinfo=timezone.utc) < now
        and t.status and t.status.value != "done"
    )
    tasks_pending = sum(1 for t in tasks_all if t.status and t.status.value == "pending")
    policies_overdue = sum(
        1 for p in policies
        if p.review_date and p.review_date.replace(tzinfo=timezone.utc) < now
    )

    EMPTY = _t("reports.mgmt_review.empty_placeholder", lang)

    _NS = "reports.mgmt_review.sections."
    sections = [
        {
            "title": _t(_NS + "a_previous_actions.title", lang),
            "content": EMPTY,
            "iso": "ISO 27001:2022 cl. 9.3.2.a",
        },
        {
            "title": _t(_NS + "b_context_changes.title", lang),
            "content": _t(
                _NS + "b_context_changes.content", lang,
                scope=(scope or EMPTY), boundaries=(boundaries or EMPTY), empty=EMPTY,
            ),
            "iso": "ISO 27001:2022 cl. 9.3.2.b",
        },
        {
            "title": _t(_NS + "c_stakeholder_needs.title", lang),
            "content": EMPTY,
            "iso": "ISO 27001:2022 cl. 9.3.2.c",
        },
        {
            "title": _t(_NS + "d_risk_management.title", lang),
            "content": _t(
                _NS + "d_risk_management.content", lang,
                total_risks=total_risks, high_risks=high_risks,
                untreated=untreated, impl_pct=impl_pct, avg_mat=avg_mat,
            ),
            "iso": "ISO 27001:2022 cl. 9.3.2.d — cl. 8.2/8.3",
        },
        {
            "title": _t(_NS + "d_incidents.title", lang),
            "content": _t(
                _NS + "d_incidents.content", lang,
                open_inc=open_inc, p1p2=p1p2,
                total_incidents=len(incidents_all),
                lessons_learned=sum(1 for i in incidents_all if i.lessons_learned),
            ),
            "iso": "ISO 27001:2022 cl. 9.3.2.d — cl. 6.1.2",
        },
        {
            "title": _t(_NS + "d_nonconformities.title", lang),
            "content": _t(_NS + "d_nonconformities.content", lang, empty=EMPTY),
            "iso": "ISO 27001:2022 cl. 9.3.2.d — cl. 10.1",
        },
        {
            "title": _t(_NS + "d_tasks.title", lang),
            "content": _t(
                _NS + "d_tasks.content", lang,
                tasks_overdue=tasks_overdue, tasks_pending=tasks_pending,
                total_tasks=len(tasks_all),
            ),
            "iso": "ISO 27001:2022 cl. 9.3.2.d — cl. 8.3",
        },
        {
            "title": _t(_NS + "d_audits.title", lang),
            "content": _t(_NS + "d_audits.content", lang, empty=EMPTY),
            "iso": "ISO 27001:2022 cl. 9.3.2.d — cl. 9.2",
        },
        {
            "title": _t(_NS + "d_policies.title", lang),
            "content": _t(
                _NS + "d_policies.content", lang,
                total_policies=len(policies), policies_overdue=policies_overdue,
                policies_draft=sum(1 for p in policies if p.status and p.status.value == 'draft'),
            ),
            "iso": "ISO 27001:2022 cl. 9.3.2.d — cl. 5.2",
        },
        {
            "title": _t(_NS + "d_suppliers.title", lang),
            "content": _t(
                _NS + "d_suppliers.content", lang,
                total_suppliers=len(suppliers),
                critical_suppliers=sum(1 for s in suppliers if s.is_critical),
                suppliers_no_assessment=sum(1 for s in suppliers if not s.last_assessment_at),
            ),
            "iso": "ISO 27001:2022 cl. 9.3.2.d — NIS2 Art. 21.2.d",
        },
        {
            "title": _t(_NS + "d_stakeholder_feedback.title", lang),
            "content": EMPTY,
            "iso": "ISO 27001:2022 cl. 9.3.2.d",
        },
        {
            "title": _t(_NS + "e_risk_results.title", lang),
            "content": _t(
                _NS + "e_risk_results.content", lang,
                total_risks=total_risks, high_risks=high_risks,
                accepted_risks=sum(1 for r in risks if r.status and r.status.value == 'accepted'),
            ),
            "iso": "ISO 27001:2022 cl. 9.3.2.e",
        },
        {
            "title": _t(_NS + "f_improvement.title", lang),
            "content": EMPTY,
            "iso": "ISO 27001:2022 cl. 9.3.2.f / cl. 10.2",
        },
        {
            "title": _t(_NS + "g_objectives.title", lang),
            "content": EMPTY,
            "iso": "ISO 27001:2022 cl. 9.3.2 / cl. 6.2",
        },
        {
            "title": _t(_NS + "h_decisions.title", lang),
            "content": EMPTY,
            "iso": "ISO 27001:2022 cl. 9.3.3",
        },
        {
            "title": _t(_NS + "i_resources.title", lang),
            "content": EMPTY,
            "iso": "ISO 27001:2022 cl. 9.3.3 / cl. 7.1",
        },
        {
            "title": _t(_NS + "j_ens.title", lang),
            "content": _t(
                _NS + "j_ens.content", lang,
                impl_pct=impl_pct, avg_mat=avg_mat, empty=EMPTY,
            ),
            "iso": "ENS RD 311/2022 — Art. 28 / Anexo II",
        },
    ]

    brand = _load_brand(db, current_user.organization_id, "management_review")
    if format == "pdf":
        return _mgmt_review_pdf(sections, org_name, scope, now_str, brand, lang)
    elif format == "excel":
        return _mgmt_review_excel(sections, org_name, scope, now_str, lang)
    else:
        return _mgmt_review_word(sections, org_name, scope, now_str, lang)


def _mgmt_review_pdf(sections, org_name, scope, now_str, brand: Optional[ReportBrand] = None, lang: str = "es"):
    """Genera el informe de revision por la direccion en PDF."""
    if brand is None:
        brand = ReportBrand()
    BRAND_PURPLE = brand.primary
    s = _styles(brand)
    _NS = "reports.mgmt_review."
    el = []
    el.append(Spacer(1, 20*mm))
    el.append(Paragraph(_t(_NS + "report_title", lang), s["TitleBrand"]))
    el.append(Paragraph(_t(_NS + "report_subtitle", lang), s["SubBrand"]))
    el.append(Spacer(1, 6))
    el.append(Paragraph(f"<b>{_t(_NS + 'label_organization', lang)}</b> {_safe(org_name)}", s["BodyBrand"]))
    el.append(Paragraph(f"<b>{_t(_NS + 'label_date', lang)}</b> {now_str}", s["BodyBrand"]))
    el.append(Paragraph(f"<b>{_t(_NS + 'label_scope', lang)}</b> {_safe(scope) or _t(_NS + 'empty_placeholder', lang)}", s["BodyBrand"]))
    el.append(PageBreak())

    for sec in sections:
        el.append(Paragraph(_safe(sec["title"]), s["H2Brand"]))
        el.append(Paragraph(f"<i>{_safe(sec['iso'])}</i>",
                             ParagraphStyle("IsoRef", parent=s["BodyBrand"],
                                            textColor=BRAND_GRAY3, fontSize=8)))
        el.append(Spacer(1, 4))
        for line in (sec["content"] or "").split("\n"):
            if line.strip():
                el.append(Paragraph(_safe(line), s["BodyBrand"]))
        el.append(Spacer(1, 12))

    # Firma
    el.append(PageBreak())
    el.append(Paragraph(_t(_NS + "approval_title", lang), s["H2Brand"]))
    firma_data = [
        [_t(_NS + "signature_role", lang), _t(_NS + "signature_name", lang),
         _t(_NS + "signature_date", lang), _t(_NS + "signature_signature", lang)],
        [_t(_NS + "role_ceo", lang), "_" * 28, "_" * 12, "_" * 20],
        [_t(_NS + "role_ciso", lang), "_" * 28, "_" * 12, "_" * 20],
        [_t(_NS + "role_auditor", lang), "_" * 28, "_" * 12, "_" * 20],
    ]
    ft = Table(firma_data, colWidths=[50*mm, 55*mm, 28*mm, 37*mm])
    ft.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), BRAND_PURPLE),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, BRAND_GRAY5]),
        ("GRID", (0, 0), (-1, -1), 0.3, BRAND_GRAY3),
        ("ROWHEIGHT", (0, 1), (-1, -1), 20*mm),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    el.append(ft)
    return _pdf_response(el, f"revision_direccion_{datetime.now().strftime('%Y%m%d')}.pdf", brand, lang)


def _mgmt_review_excel(sections, org_name, scope, now_str, lang: str = "es"):
    """Genera el informe de revision por la direccion en Excel."""
    import openpyxl
    from openpyxl.styles import Alignment, Font, PatternFill, Border, Side

    _NS = "reports.mgmt_review."
    empty_ph = _t(_NS + "empty_placeholder", lang)

    wb = openpyxl.Workbook()
    thin = Side(style="thin", color="D1D5DB")
    brd = Border(left=thin, right=thin, top=thin, bottom=thin)
    hdr_fill = PatternFill("solid", fgColor="59008D")
    emp_fill = PatternFill("solid", fgColor="FFFBEB")  # amarillo claro = a rellenar

    ws = wb.active
    ws.title = _t(_NS + "excel_sheet_title", lang)
    ws.column_dimensions["A"].width = 45
    ws.column_dimensions["B"].width = 80

    ws.append([_t(_NS + "report_title", lang)])
    ws["A1"].font = Font(size=14, bold=True, color="59008D")
    ws.append([_t(_NS + "field_organization", lang, org=org_name, date=now_str)])
    ws.append([_t(_NS + "field_scope", lang, scope=(scope or empty_ph))])
    ws.append([_t(_NS + "excel_header_norma", lang), _t(_NS + "standard_reference", lang)])
    ws.append([])

    for sec in sections:
        row_hdr = ws.max_row + 1
        ws.append([sec["title"], sec["iso"]])
        for col in [1, 2]:
            c = ws.cell(row=row_hdr, column=col)
            c.fill = hdr_fill
            c.font = Font(bold=True, color="FFFFFF", size=9)
            c.border = brd

        content = sec["content"] or ""
        needs_fill = empty_ph in content
        row_cont = ws.max_row + 1
        ws.append([_t(_NS + "excel_content_label", lang), content])
        for col in [1, 2]:
            c = ws.cell(row=row_cont, column=col)
            c.alignment = Alignment(wrap_text=True, vertical="top")
            c.border = brd
            if needs_fill and col == 2:
                c.fill = emp_fill
                c.font = Font(italic=True, color="92400E")
        ws.row_dimensions[row_cont].height = max(30, content.count("\n") * 14 + 20)
        ws.append([])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename=revision_direccion_{datetime.now().strftime('%Y%m%d')}.xlsx"},
    )


def _mgmt_review_word(sections, org_name, scope, now_str, lang: str = "es"):
    """Genera el informe de revision por la direccion en Word (.docx)."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(500, _t("reports.mgmt_review.docx_not_installed", lang))

    _NS = "reports.mgmt_review."
    empty_ph = _t(_NS + "empty_placeholder", lang)

    doc = Document()

    # Estilos
    style_title = doc.styles["Title"]
    style_title.font.size = Pt(22)
    style_title.font.color.rgb = RGBColor(0x59, 0x00, 0x8D)

    doc.add_heading(_t(_NS + "report_title", lang), 0)
    doc.add_paragraph(f"{_t(_NS + 'label_organization', lang)} {org_name}")
    doc.add_paragraph(f"{_t(_NS + 'label_date', lang)} {now_str}")
    doc.add_paragraph(f"{_t(_NS + 'label_scope', lang)} {scope or empty_ph}")
    doc.add_paragraph(_t(_NS + "standard_reference_word", lang))
    doc.add_page_break()

    for sec in sections:
        doc.add_heading(sec["title"], 2)
        iso_p = doc.add_paragraph(sec["iso"])
        iso_p.runs[0].italic = True
        iso_p.runs[0].font.size = Pt(8)
        iso_p.runs[0].font.color.rgb = RGBColor(0x9D, 0x9D, 0x9D)

        content = sec["content"] or empty_ph
        for line in content.split("\n"):
            if line.strip():
                p = doc.add_paragraph(line.strip())
                if empty_ph in line:
                    for run in p.runs:
                        run.italic = True
                        run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)

        doc.add_paragraph()

    # Seccion de firmas
    doc.add_page_break()
    doc.add_heading(_t(_NS + "approval_title", lang), 2)
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    for i, hdr in enumerate([
        _t(_NS + "signature_role", lang), _t(_NS + "signature_name", lang),
        _t(_NS + "signature_date", lang), _t(_NS + "signature_signature", lang),
    ]):
        table.cell(0, i).text = hdr
        table.cell(0, i).paragraphs[0].runs[0].bold = True
    for row_idx, cargo in enumerate(
        [_t(_NS + "role_ceo", lang), _t(_NS + "role_ciso_word", lang), _t(_NS + "role_auditor", lang)], 1
    ):
        table.cell(row_idx, 0).text = cargo
        for col in [1, 2, 3]:
            table.cell(row_idx, col).text = ""

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=revision_direccion_{datetime.now().strftime('%Y%m%d')}.docx"},
    )


# ============================================================
# Informe TPRM (Third-Party Risk Management)
# ============================================================

_TPRM_LEVEL_COLORS = {
    "critical": colors.HexColor("#FCA5A5"),
    "high": colors.HexColor("#FED7AA"),
    "medium": colors.HexColor("#FEF9C3"),
    "low": colors.HexColor("#DCFCE7"),
}


@router.get("/tprm")
def tprm_report(request: Request, db: Session = Depends(get_db),
                current_user: User = Depends(get_current_user)):
    """Informe consolidado de riesgo de terceros (TPRM): inventario por tier y
    riesgo, evaluaciones vigentes, hallazgos abiertos con SLA y estado de
    cuestionarios. Un unico PDF para direccion / auditoria de la cadena de
    suministro (ISO 27001 A.5.19-A.5.22, ISO 27005)."""
    from app.models import (
        Supplier, SupplierQuestionnaire, VendorIssue, VendorIssueStatus,
        VendorRiskAssessment,
    )
    lang = get_lang(request)
    brand = _load_brand(db, current_user.organization_id, "tprm")
    BRAND_PURPLE = brand.primary
    BRAND_ORANGE = brand.secondary
    s = _styles(brand)
    now = datetime.now(timezone.utc)
    el = []

    el.append(Paragraph(_t("reports.tprm.title", lang), s["TitleBrand"]))
    el.append(Paragraph(_t("reports.tprm.subtitle", lang), s["SubBrand"]))
    ctx = filter_by_org(db.query(RiskContext), RiskContext, current_user).first()
    org_name = (ctx.organization_name if ctx else None) or "-"
    el.append(Paragraph(f"<b>{_t('reports.tprm.organization_label', lang)}</b> {_safe(org_name)}", s["BodyBrand"]))
    el.append(Paragraph(f"<b>{_t('reports.tprm.date_label', lang)}</b> {now.strftime('%d/%m/%Y')}", s["BodyBrand"]))
    el.append(Spacer(1, 12))

    suppliers = filter_by_org(db.query(Supplier), Supplier, current_user).all()
    total = len(suppliers)
    criticos = sum(1 for x in suppliers if x.is_critical)

    # ── Resumen ejecutivo ────────────────────────────────────────────────
    el.append(Paragraph(_t("reports.tprm.summary_title", lang), s["H2Brand"]))
    by_tier: dict[str, int] = {}
    for sup in suppliers:
        key = sup.tier.value if sup.tier else "-"
        by_tier[key] = by_tier.get(key, 0) + 1
    by_level: dict[str, int] = {}
    for sup in suppliers:
        key = (sup.risk_level or "medium").lower()
        by_level[key] = by_level.get(key, 0) + 1
    el.append(Paragraph(
        _t("reports.tprm.summary_body", lang, total=total, critical=criticos),
        s["BodyBrand"]))
    tier_txt = ", ".join(f"{k}: {v}" for k, v in sorted(by_tier.items())) or "-"
    el.append(Paragraph(f"<b>{_t('reports.tprm.by_tier_label', lang)}</b> {tier_txt}", s["BodyBrand"]))
    el.append(Spacer(1, 12))

    # ── Inventario de proveedores ────────────────────────────────────────
    el.append(Paragraph(_t("reports.tprm.inventory_title", lang), s["H2Brand"]))
    data = [[
        _t("reports.tprm.h_code", lang), _t("reports.tprm.h_name", lang),
        _t("reports.tprm.h_tier", lang), _t("reports.tprm.h_critical", lang),
        _t("reports.tprm.h_inherent", lang), _t("reports.tprm.h_residual", lang),
        _t("reports.tprm.h_next", lang),
    ]]
    for sup in sorted(suppliers, key=lambda x: (x.residual_risk_score or 0), reverse=True):
        data.append([
            sup.code or "-", _safe(sup.name, 34),
            sup.tier.value if sup.tier else "-",
            _t("reports.tprm.yes", lang) if sup.is_critical else "-",
            str(sup.inherent_risk_score) if sup.inherent_risk_score is not None else "-",
            str(sup.residual_risk_score) if sup.residual_risk_score is not None else "-",
            sup.next_assessment_at.strftime("%d/%m/%y") if sup.next_assessment_at else "-",
        ])
    if len(data) > 1:
        tbl = Table(data, repeatRows=1, colWidths=[18*mm, 46*mm, 20*mm, 16*mm, 22*mm, 22*mm, 24*mm])
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), BRAND_PURPLE),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, BRAND_GRAY5]),
            ("GRID", (0, 0), (-1, -1), 0.25, BRAND_GRAY3),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        el.append(tbl)
    else:
        el.append(Paragraph(_t("reports.tprm.no_suppliers", lang), s["BodyBrand"]))

    # ── Evaluaciones vigentes ────────────────────────────────────────────
    assessments = filter_by_org(
        db.query(VendorRiskAssessment), VendorRiskAssessment, current_user
    ).filter(VendorRiskAssessment.is_current == True).all()  # noqa: E712
    if assessments:
        el.append(Spacer(1, 12))
        el.append(Paragraph(_t("reports.tprm.assessments_title", lang), s["H2Brand"]))
        adata = [[
            _t("reports.tprm.h_code", lang), _t("reports.tprm.h_supplier", lang),
            _t("reports.tprm.h_period", lang), _t("reports.tprm.h_residual_level", lang),
            _t("reports.tprm.h_recommendation", lang), _t("reports.tprm.h_valid_until", lang),
        ]]
        for a in assessments:
            adata.append([
                a.code or "-", _safe(a.supplier_name, 30),
                a.period_label or "-",
                (a.residual_risk_level or "-"),
                a.recommendation.value if a.recommendation else "-",
                a.valid_until.strftime("%d/%m/%y") if a.valid_until else "-",
            ])
        atbl = Table(adata, repeatRows=1, colWidths=[20*mm, 40*mm, 22*mm, 24*mm, 34*mm, 24*mm])
        atbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), BRAND_PURPLE),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, BRAND_GRAY5]),
            ("GRID", (0, 0), (-1, -1), 0.25, BRAND_GRAY3),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        el.append(atbl)

    # ── Hallazgos abiertos con SLA ───────────────────────────────────────
    issues = filter_by_org(db.query(VendorIssue), VendorIssue, current_user).filter(
        VendorIssue.status != VendorIssueStatus.CLOSED
    ).order_by(VendorIssue.severity.desc()).all()
    el.append(Spacer(1, 12))
    el.append(Paragraph(_t("reports.tprm.issues_title", lang), s["H2Brand"]))
    if issues:
        idata = [[
            _t("reports.tprm.h_code", lang), _t("reports.tprm.h_supplier", lang),
            _t("reports.tprm.h_issue", lang), _t("reports.tprm.h_severity", lang),
            _t("reports.tprm.h_status", lang), _t("reports.tprm.h_due", lang),
        ]]
        for i in issues:
            overdue = bool(i.due_date and i.due_date.replace(tzinfo=timezone.utc) < now)
            due_txt = (i.due_date.strftime("%d/%m/%y") if i.due_date else "-")
            if overdue:
                due_txt = _t("reports.tprm.overdue", lang, date=due_txt)
            idata.append([
                i.code or "-", _safe(i.supplier_name, 26), _safe(i.title, 40),
                i.severity.value if i.severity else "-",
                i.status.value if i.status else "-", due_txt,
            ])
        itbl = Table(idata, repeatRows=1, colWidths=[20*mm, 34*mm, 50*mm, 20*mm, 20*mm, 24*mm])
        itbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), BRAND_ORANGE),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, BRAND_GRAY5]),
            ("GRID", (0, 0), (-1, -1), 0.25, BRAND_GRAY3),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        el.append(itbl)
    else:
        el.append(Paragraph(_t("reports.tprm.no_issues", lang), s["BodyBrand"]))

    # ── Estado de cuestionarios ──────────────────────────────────────────
    questionnaires = filter_by_org(
        db.query(SupplierQuestionnaire), SupplierQuestionnaire, current_user
    ).all()
    if questionnaires:
        pend = sum(1 for q in questionnaires if not q.submitted_at)
        done = len(questionnaires) - pend
        el.append(Spacer(1, 12))
        el.append(Paragraph(_t("reports.tprm.questionnaires_title", lang), s["H2Brand"]))
        el.append(Paragraph(
            _t("reports.tprm.questionnaires_body", lang,
               total=len(questionnaires), done=done, pending=pend),
            s["BodyBrand"]))

    return _pdf_response(el, "informe_tprm.pdf", brand, lang)


# ============================================================
# Informes generados por IA
# ============================================================

REPORT_LABEL = {
    "treatment_plan": "Plan de Tratamiento de Riesgos",
    "executive_dashboard": "Dashboard Ejecutivo",
    "committee_minutes": "Acta de Comite de Seguridad",
    "followup_report": "Informe de Seguimiento ISO 27005",
}


def _report_label(report_type: str, lang: str = "es") -> str:
    """Etiqueta visible del tipo de informe IA, en el idioma de UI."""
    if report_type not in REPORT_LABEL:
        return _t("reports.ai_report.label_fallback", lang)
    return _t(f"reports.ai_report.label_{report_type}", lang)


# ============================================================
# Informe del Estado del SGSI — multi-modulo, sin IA
# ============================================================

def _kpi_table(data_rows: list[tuple], s, brand: Optional[ReportBrand] = None, lang: str = "es") -> "Table":
    """Tabla de 2 columnas: etiqueta | valor."""
    if brand is None:
        brand = ReportBrand()
    rows = [[_t("reports.sgsi_status.kpi_header_indicator", lang), _t("reports.sgsi_status.kpi_header_value", lang)]] + list(data_rows)
    t = Table(rows, colWidths=[110 * mm, 60 * mm])
    style = [
        ("BACKGROUND", (0, 0), (-1, 0), brand.primary),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), brand.font_bold),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, BRAND_GRAY5]),
        ("GRID", (0, 0), (-1, -1), 0.25, BRAND_GRAY3),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]
    t.setStyle(TableStyle(style))
    return t


@router.get("/sgsi-status")
def sgsi_status_report(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Informe del Estado del SGSI — resumen multi-modulo sin IA. Descarga instantanea."""
    lang = get_lang(request)
    from datetime import timezone
    brand = _load_brand(db, current_user.organization_id, "sgsi_status")
    BRAND_PURPLE = brand.primary
    BRAND_ORANGE = brand.secondary
    s = _styles(brand)
    el = []

    ctx = filter_by_org(db.query(RiskContext), RiskContext, current_user).first()
    org = ctx.organization_name if ctx else _t("reports.sgsi_status.default_org_name", lang)
    now = datetime.now(timezone.utc)

    # --- Portada ---
    el.append(Spacer(1, 20))
    el.append(Paragraph(_t("reports.sgsi_status.cover_title", lang), s["TitleBrand"]))
    el.append(Paragraph(_t("reports.sgsi_status.cover_subtitle", lang), s["SubBrand"]))
    el.append(Spacer(1, 8))
    el.append(Paragraph(f"<b>{_t('reports.sgsi_status.org_label', lang)}</b> {org}", s["BodyBrand"]))
    el.append(Paragraph(f"<b>{_t('reports.sgsi_status.date_label', lang)}</b> {now.strftime('%d/%m/%Y %H:%M UTC')}", s["BodyBrand"]))
    if ctx:
        el.append(Paragraph(f"<b>{_t('reports.sgsi_status.scope_label', lang)}</b> {ctx.scope or '-'}", s["BodyBrand"]))
        el.append(Paragraph(f"<b>{_t('reports.sgsi_status.appetite_label', lang)}</b> {_t('reports.sgsi_status.appetite_level', lang, level=ctx.risk_appetite)}", s["BodyBrand"]))
    el.append(Spacer(1, 12))

    # --- Seccion 1: Riesgos ---
    risks = filter_by_org(db.query(Risk), Risk, current_user).all()
    total_r = len(risks)
    high_r = sum(1 for r in risks if (r.residual_level or 0) >= 5)
    medium_r = sum(1 for r in risks if 3 <= (r.residual_level or 0) < 5)
    low_r = sum(1 for r in risks if (r.residual_level or 0) < 3)
    treated_r = sum(1 for r in risks if r.treatment_option is not None)
    overdue_r = sum(
        1 for r in risks
        if r.treatment_due_date and r.status not in (RiskStatus.ACCEPTED, RiskStatus.CLOSED)
        and r.treatment_due_date.replace(tzinfo=timezone.utc) < now
    )

    el.append(Paragraph(_t("reports.sgsi_status.section1_title", lang), s["H2Brand"]))
    el.append(_kpi_table([
        (_t("reports.sgsi_status.kpi_total_risks", lang), str(total_r)),
        (_t("reports.sgsi_status.kpi_high_risks", lang), str(high_r)),
        (_t("reports.sgsi_status.kpi_medium_risks", lang), str(medium_r)),
        (_t("reports.sgsi_status.kpi_low_risks", lang), str(low_r)),
        (_t("reports.sgsi_status.kpi_treated_risks", lang), str(treated_r)),
        (_t("reports.sgsi_status.kpi_overdue_treatments", lang), str(overdue_r)),
    ], s, brand, lang))
    el.append(Spacer(1, 8))

    # Top riesgos altos
    top_high = sorted(
        [r for r in risks if (r.residual_level or 0) >= 5],
        key=lambda r: -(r.residual_level or 0)
    )[:10]
    if top_high:
        el.append(Paragraph(_t("reports.sgsi_status.top_high_risks_title", lang), s["BodyBrand"]))
        el.append(Spacer(1, 4))
        th_data = [[
            _t("reports.sgsi_status.risk_table_header_code", lang), _t("reports.sgsi_status.risk_table_header_asset", lang),
            _t("reports.sgsi_status.risk_table_header_threat", lang), _t("reports.sgsi_status.risk_table_header_residual_level", lang),
            _t("reports.sgsi_status.risk_table_header_treatment", lang),
        ]]
        for r in top_high:
            th_data.append([
                r.code,
                (r.asset.name[:30] if r.asset else "-"),
                (r.threat.name[:35] if r.threat else "-"),
                str(r.residual_level or 0),
                (r.treatment_option.value if r.treatment_option else _t("reports.sgsi_status.not_defined", lang)),
            ])
        th = Table(th_data, repeatRows=1, colWidths=[18*mm, 42*mm, 50*mm, 18*mm, 32*mm])
        th.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), BRAND_ORANGE),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#FEE2E2"), colors.HexColor("#FEF2F2")]),
            ("GRID", (0, 0), (-1, -1), 0.25, BRAND_GRAY3),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        el.append(th)
    el.append(Spacer(1, 10))

    # --- Seccion 2: Controles ---
    el.append(Paragraph(_t("reports.sgsi_status.section2_title", lang), s["H2Brand"]))
    impls = filter_by_org(db.query(ControlImplementation), ControlImplementation, current_user).all()
    total_impl = len(impls)
    impl_done = sum(1 for i in impls if i.status and i.status.value == "implemented")
    impl_partial = sum(1 for i in impls if i.status and i.status.value == "partial")
    impl_planned = sum(1 for i in impls if i.status and i.status.value == "planned")
    avg_mat = (sum(i.maturity or 0 for i in impls) / total_impl) if total_impl else 0
    overdue_ctrl = sum(
        1 for i in impls
        if i.next_review and i.status and i.status.value != "not_implemented"
        and i.next_review.replace(tzinfo=timezone.utc) < now
    )
    el.append(_kpi_table([
        (_t("reports.sgsi_status.kpi_active_controls", lang), str(total_impl)),
        (_t("reports.sgsi_status.kpi_status_implemented", lang), str(impl_done)),
        (_t("reports.sgsi_status.kpi_status_partial", lang), str(impl_partial)),
        (_t("reports.sgsi_status.kpi_status_planned", lang), str(impl_planned)),
        (_t("reports.sgsi_status.kpi_avg_maturity", lang), f"{avg_mat:.1f}"),
        (_t("reports.sgsi_status.kpi_overdue_reviews", lang), str(overdue_ctrl)),
    ], s, brand, lang))
    el.append(Spacer(1, 10))

    # --- Seccion 3: Incidentes ---
    el.append(Paragraph(_t("reports.sgsi_status.section3_title", lang), s["H2Brand"]))
    incidents = filter_by_org(db.query(Incident), Incident, current_user).all()
    total_inc = len(incidents)
    open_inc = sum(1 for i in incidents if i.status != IncidentStatus.CLOSED)
    p1p2 = sum(1 for i in incidents if i.severity in ("p1", "p2") and i.status != IncidentStatus.CLOSED)
    nis2_pending = sum(
        1 for i in incidents
        if i.nis2_notification_required and not i.nis2_notification_sent_at
        and i.status != IncidentStatus.CLOSED
    )
    el.append(_kpi_table([
        (_t("reports.sgsi_status.kpi_total_incidents", lang), str(total_inc)),
        (_t("reports.sgsi_status.kpi_open_incidents", lang), str(open_inc)),
        (_t("reports.sgsi_status.kpi_critical_incidents", lang), str(p1p2)),
        (_t("reports.sgsi_status.kpi_nis2_pending", lang), str(nis2_pending)),
    ], s, brand, lang))
    el.append(Spacer(1, 10))

    # --- Seccion 4: Tareas de tratamiento ---
    el.append(Paragraph(_t("reports.sgsi_status.section4_title", lang), s["H2Brand"]))
    tasks = filter_by_org(db.query(TreatmentTask), TreatmentTask, current_user).all()
    total_t = len(tasks)
    done_t = sum(1 for t in tasks if t.status == TaskStatus.DONE)
    inprog_t = sum(1 for t in tasks if t.status == TaskStatus.IN_PROGRESS)
    pend_t = sum(1 for t in tasks if t.status == TaskStatus.PENDING)
    overdue_t = sum(
        1 for t in tasks
        if t.due_date and t.status != TaskStatus.DONE
        and t.due_date.replace(tzinfo=timezone.utc) < now
    )
    el.append(_kpi_table([
        (_t("reports.sgsi_status.kpi_total_tasks", lang), str(total_t)),
        (_t("reports.sgsi_status.kpi_done_tasks", lang), str(done_t)),
        (_t("reports.sgsi_status.kpi_inprogress_tasks", lang), str(inprog_t)),
        (_t("reports.sgsi_status.kpi_pending_tasks", lang), str(pend_t)),
        (_t("reports.sgsi_status.kpi_overdue_tasks", lang), str(overdue_t)),
    ], s, brand, lang))
    el.append(Spacer(1, 10))

    # --- Seccion 5: Politicas ---
    el.append(Paragraph(_t("reports.sgsi_status.section5_title", lang), s["H2Brand"]))
    policies = filter_by_org(db.query(Policy), Policy, current_user).all()
    total_pol = len(policies)
    pub_pol = sum(1 for p in policies if p.status == PolicyStatus.PUBLISHED)
    rev_pol = sum(1 for p in policies if p.status == PolicyStatus.REVIEW)
    draft_pol = sum(1 for p in policies if p.status == PolicyStatus.DRAFT)
    overdue_pol = sum(
        1 for p in policies
        if p.review_date and p.status != PolicyStatus.OBSOLETE
        and p.review_date.replace(tzinfo=timezone.utc) < now
    )
    el.append(_kpi_table([
        (_t("reports.sgsi_status.kpi_total_policies", lang), str(total_pol)),
        (_t("reports.sgsi_status.kpi_published_policies", lang), str(pub_pol)),
        (_t("reports.sgsi_status.kpi_review_policies", lang), str(rev_pol)),
        (_t("reports.sgsi_status.kpi_draft_policies", lang), str(draft_pol)),
        (_t("reports.sgsi_status.kpi_overdue_policy_review", lang), str(overdue_pol)),
    ], s, brand, lang))
    el.append(Spacer(1, 10))

    # --- Seccion 6: RGPD ---
    el.append(Paragraph(_t("reports.sgsi_status.section6_title", lang), s["H2Brand"]))
    activities = filter_by_org(db.query(ProcessingActivity), ProcessingActivity, current_user).all()
    act_ids = [a.id for a in activities]
    dpias = db.query(DPIA).filter(DPIA.activity_id.in_(act_ids)).all() if act_ids else []
    total_act = len(activities)
    req_dpia = sum(1 for a in activities if a.requires_dpia)
    eu_transfer = sum(1 for a in activities if a.transfers_outside_eu)
    total_dp = len(dpias)
    pend_dp = sum(1 for d in dpias if d.status == DPIAStatus.PENDING)
    el.append(_kpi_table([
        (_t("reports.sgsi_status.kpi_total_activities", lang), str(total_act)),
        (_t("reports.sgsi_status.kpi_dpia_required", lang), str(req_dpia)),
        (_t("reports.sgsi_status.kpi_eu_transfers", lang), str(eu_transfer)),
        (_t("reports.sgsi_status.kpi_total_dpias", lang), str(total_dp)),
        (_t("reports.sgsi_status.kpi_pending_dpias", lang), str(pend_dp)),
    ], s, brand, lang))
    el.append(Spacer(1, 10))

    # --- Seccion 7: Nota metodologica ---
    el.append(PageBreak())
    el.append(Paragraph(_t("reports.sgsi_status.section7_title", lang), s["H2Brand"]))
    el.append(Paragraph(
        _t("reports.sgsi_status.methodology_note", lang),
        s["BodyBrand"],
    ))
    el.append(Spacer(1, 8))
    refs = [
        [_t("reports.sgsi_status.refs_table_header_standard", lang), _t("reports.sgsi_status.refs_table_header_coverage", lang)],
        ["ISO/IEC 27005:2018", _t("reports.sgsi_status.ref_iso27005_coverage", lang)],
        ["ISO/IEC 27002:2022", _t("reports.sgsi_status.ref_iso27002_coverage", lang)],
        ["ISO/IEC 27001:2022", _t("reports.sgsi_status.ref_iso27001_coverage", lang)],
        ["NIS2 (Directiva UE 2022/2555)", _t("reports.sgsi_status.ref_nis2_coverage", lang)],
        ["RGPD / GDPR (Reglamento UE 2016/679)", _t("reports.sgsi_status.ref_gdpr_coverage", lang)],
    ]
    ref_t = Table(refs, repeatRows=1, colWidths=[70 * mm, 100 * mm])
    ref_t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), BRAND_PURPLE),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, BRAND_GRAY5]),
        ("GRID", (0, 0), (-1, -1), 0.25, BRAND_GRAY3),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    el.append(ref_t)

    filename = f"sgsi_status_{now.strftime('%Y%m%d')}.pdf"
    return _pdf_response(el, filename, brand, lang)


class AiReportIn(BaseModel):
    report_type: str
    format: str = "pdf"  # pdf | excel


@router.post("/ai-generate")
def ai_generate(body: AiReportIn, request: Request, db: Session = Depends(get_db),
                current_user: User = Depends(get_current_user)):
    """Genera un informe ejecutivo usando Claude API y lo devuelve como PDF o Excel."""
    lang = get_lang(request)
    if body.report_type not in REPORT_LABEL:
        raise HTTPException(422, _t("reports.ai_report.invalid_report_type", lang, options=list(REPORT_LABEL)))
    if body.format not in ("pdf", "excel"):
        raise HTTPException(422, _t("reports.ai_report.invalid_format", lang))

    # Resolver API key del tenant (configurada en IA -> Configuracion)
    from app.routers.ai_config import resolve_api_key
    ai_cfg = filter_by_org(db.query(AiConfig), AiConfig, current_user).first()
    api_key = resolve_api_key(ai_cfg)

    try:
        # M6/C1: propagar org_id para que _collect() filtre por tenant
        content = report_ai_service.generate(
            body.report_type, db, api_key=api_key,
            org_id=current_user.organization_id,
            lang=lang,
        )
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(502, _t("reports.generation_failed", lang))

    brand = _load_brand(db, current_user.organization_id, body.report_type)
    if body.format == "excel":
        return _ai_report_excel(content, body.report_type, lang)
    return _ai_report_pdf(content, body.report_type, brand, lang)


def _safe(text, max_chars=None):
    """Convierte a string y opcionalmente trunca para ReportLab."""
    if not text:
        return ""
    s = str(text)
    if max_chars:
        s = s[:max_chars]
    # ReportLab no soporta algunos caracteres de control
    return s.replace("\x00", "").replace("\r", "")


def _ai_report_pdf(content: dict, report_type: str, brand: Optional[ReportBrand] = None, lang: str = "es"):
    """Convierte el JSON de Claude a un PDF con ReportLab."""
    if brand is None:
        brand = ReportBrand()
    BRAND_PURPLE = brand.primary
    BRAND_ORANGE = brand.secondary
    s = _styles(brand)
    el = []
    label = _report_label(report_type, lang)

    # Portada
    el.append(Spacer(1, 30 * mm))
    el.append(Paragraph(_safe(content.get("title", label)), s["TitleBrand"]))
    org = _safe(content.get("organization", ""))
    if org:
        el.append(Paragraph(org, s["SubBrand"]))
    date_str = _safe(content.get("date", datetime.now().strftime("%d/%m/%Y")))
    el.append(Paragraph(_t("reports.ai_report.date_label", lang, date=date_str), s["BodyBrand"]))
    el.append(Spacer(1, 8))

    def add_section(title, text):
        if text:
            el.append(Paragraph(title, s["H2Brand"]))
            el.append(Paragraph(_safe(text), s["BodyBrand"]))
            el.append(Spacer(1, 6))

    def add_list(title, items):
        if items:
            el.append(Paragraph(title, s["H2Brand"]))
            for item in items:
                el.append(Paragraph(f"• {_safe(item)}", s["BodyBrand"]))
            el.append(Spacer(1, 6))

    # Secciones comunes
    add_section(_t("reports.ai_report.executive_summary", lang), content.get("executive_summary"))

    # Secciones especificas por tipo
    if report_type == "treatment_plan":
        add_section(_t("reports.ai_report.risk_appetite_analysis", lang), content.get("risk_appetite_analysis"))
        risks = content.get("risks", [])
        if risks:
            el.append(Paragraph(_t("reports.ai_report.treatment_plans_by_risk", lang), s["H2Brand"]))
            el.append(Spacer(1, 4))
            data = [[
                _t("reports.ai_report.header_code", lang), _t("reports.ai_report.header_priority", lang),
                _t("reports.ai_report.header_effort", lang), _t("reports.ai_report.header_target_level", lang),
                _t("reports.ai_report.header_actions", lang),
            ]]
            for r in risks:
                actions = "\n".join(f"• {a}" for a in (r.get("recommended_actions") or [])[:3])
                data.append([
                    _safe(r.get("code", "")),
                    _safe(r.get("priority", "")),
                    _safe(r.get("estimated_effort", "")),
                    str(r.get("target_residual_level", "")),
                    _safe(actions, 200),
                ])
            t = Table(data, repeatRows=1, colWidths=[22*mm, 28*mm, 20*mm, 22*mm, 73*mm])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), BRAND_PURPLE),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, BRAND_GRAY5]),
                ("GRID", (0, 0), (-1, -1), 0.25, BRAND_GRAY3),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            el.append(t)
            el.append(Spacer(1, 8))
            # Narrativas detalladas
            for r in risks:
                el.append(Paragraph(f"<b>{_safe(r.get('code',''))}</b> — {_safe(r.get('treatment_narrative',''))}",
                                    s["BodyBrand"]))
                if r.get("success_metrics"):
                    el.append(Paragraph(f"<i>{_t('reports.ai_report.success_metrics_label', lang)}</i> {_safe(r['success_metrics'])}", s["BodyBrand"]))
                el.append(Spacer(1, 4))
        add_section(_t("reports.ai_report.implementation_roadmap", lang), content.get("implementation_roadmap"))
        add_section(_t("reports.ai_report.conclusion", lang), content.get("conclusion"))

    elif report_type == "executive_dashboard":
        add_list(_t("reports.ai_report.key_findings", lang), content.get("key_findings"))
        add_section(_t("reports.ai_report.risk_posture", lang), content.get("risk_posture_explanation"))
        add_section(_t("reports.ai_report.critical_risks", lang), content.get("top_risks_narrative"))
        add_section(_t("reports.ai_report.control_effectiveness", lang), content.get("control_effectiveness"))
        add_section(_t("reports.ai_report.compliance_status", lang), content.get("compliance_status"))
        add_list(_t("reports.ai_report.critical_actions_required", lang), content.get("critical_actions"))
        add_section(_t("reports.ai_report.kpi_analysis", lang), content.get("kpi_commentary"))
        add_section(_t("reports.ai_report.next_review_recommendation", lang), content.get("next_review_recommendation"))

        # Tabla de estadisticas
        meta = content.get("_meta", {})
        stats = meta.get("stats", {})
        if stats:
            el.append(Paragraph(_t("reports.ai_report.risk_register_stats", lang), s["H2Brand"]))
            stat_data = [
                [_t("reports.ai_report.stat_total", lang), str(stats.get("total", 0))],
                [_t("reports.ai_report.stat_critical", lang), str(stats.get("critical", 0))],
                [_t("reports.ai_report.stat_high", lang), str(stats.get("high", 0))],
                [_t("reports.ai_report.stat_medium", lang), str(stats.get("medium", 0))],
                [_t("reports.ai_report.stat_low", lang), str(stats.get("low", 0))],
                [_t("reports.ai_report.stat_with_treatment", lang), str(stats.get("with_treatment", 0))],
                [_t("reports.ai_report.stat_accepted", lang), str(stats.get("accepted", 0))],
            ]
            t = Table(stat_data, colWidths=[60*mm, 30*mm])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (0, -1), BRAND_GRAY5),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.25, BRAND_GRAY3),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ALIGN", (1, 0), (1, -1), "CENTER"),
            ]))
            el.append(t)

    elif report_type == "committee_minutes":
        session_info = content.get("session_number", "")
        if session_info:
            add_section(_t("reports.ai_report.session_number", lang), session_info)
        add_section(_t("reports.ai_report.attendees_note", lang), content.get("attendees_note"))
        add_list(_t("reports.ai_report.agenda", lang), content.get("agenda"))
        add_section(_t("reports.ai_report.risk_register_review", lang), content.get("risk_register_review"))
        # Riesgos aceptados
        accepted = content.get("accepted_risks", [])
        if accepted:
            el.append(Paragraph(_t("reports.ai_report.accepted_risks", lang), s["H2Brand"]))
            for ar in accepted:
                el.append(Paragraph(
                    f"<b>{_safe(ar.get('code',''))}</b>: {_safe(ar.get('rationale',''))}",
                    s["BodyBrand"],
                ))
            el.append(Spacer(1, 6))
        add_section(_t("reports.ai_report.treatment_followup", lang), content.get("treatment_followup"))
        add_list(_t("reports.ai_report.decisions_adopted", lang), content.get("decisions"))
        # Acciones
        actions = content.get("action_items", [])
        if actions:
            el.append(Paragraph(_t("reports.ai_report.agreed_actions", lang), s["H2Brand"]))
            data = [[_t("reports.ai_report.header_action", lang), _t("reports.ai_report.header_responsible", lang), _t("reports.ai_report.header_deadline", lang)]]
            for a in actions:
                data.append([_safe(a.get("action", ""), 80),
                              _safe(a.get("responsible", "")),
                              _safe(a.get("deadline", ""))])
            t = Table(data, repeatRows=1, colWidths=[90*mm, 45*mm, 30*mm])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), BRAND_PURPLE),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.25, BRAND_GRAY3),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            el.append(t)
        add_list(_t("reports.ai_report.next_session_topics", lang), content.get("next_session_topics"))
        add_section(_t("reports.ai_report.closing", lang), content.get("closing_note"))

    elif report_type == "followup_report":
        add_section(_t("reports.ai_report.subtitle", lang), content.get("subtitle"))
        cl12 = content.get("cl12_assessment", {})
        if cl12:
            el.append(Paragraph(_t("reports.ai_report.cl12_assessment", lang), s["H2Brand"]))
            for key, label_text in [
                ("monitoring_adequacy", _t("reports.ai_report.cl12_monitoring_adequacy", lang)),
                ("review_frequency", _t("reports.ai_report.cl12_review_frequency", lang)),
                ("improvement_actions", _t("reports.ai_report.cl12_improvement_actions", lang)),
                ("context_changes", _t("reports.ai_report.cl12_context_changes", lang)),
            ]:
                if cl12.get(key):
                    el.append(Paragraph(f"<b>{label_text}:</b> {_safe(cl12[key])}", s["BodyBrand"]))
                    el.append(Spacer(1, 4))
        kpi = content.get("kpi_analysis", {})
        if kpi:
            el.append(Paragraph(_t("reports.ai_report.kpi_analysis", lang), s["H2Brand"]))
            for key, label_text in [
                ("risk_reduction_trend", _t("reports.ai_report.kpi_risk_reduction_trend", lang)),
                ("treatment_effectiveness", _t("reports.ai_report.kpi_treatment_effectiveness", lang)),
                ("control_coverage", _t("reports.ai_report.kpi_control_coverage", lang)),
                ("pending_actions", _t("reports.ai_report.kpi_pending_actions", lang)),
            ]:
                if kpi.get(key):
                    el.append(Paragraph(f"<b>{label_text}:</b> {_safe(kpi[key])}", s["BodyBrand"]))
                    el.append(Spacer(1, 4))
        add_list(_t("reports.ai_report.strengths_identified", lang), content.get("strengths"))
        add_list(_t("reports.ai_report.weaknesses_identified", lang), content.get("weaknesses"))
        recs = content.get("recommendations", [])
        if recs:
            el.append(Paragraph(_t("reports.ai_report.recommendations", lang), s["H2Brand"]))
            data = [[_t("reports.ai_report.header_area", lang), _t("reports.ai_report.header_recommendation", lang), _t("reports.ai_report.header_priority", lang), _t("reports.ai_report.header_iso_ref", lang)]]
            for rec in recs:
                data.append([
                    _safe(rec.get("area", ""), 30),
                    _safe(rec.get("recommendation", ""), 100),
                    _safe(rec.get("priority", "")),
                    _safe(rec.get("iso_reference", "")),
                ])
            t = Table(data, repeatRows=1, colWidths=[30*mm, 90*mm, 18*mm, 25*mm])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), BRAND_PURPLE),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.25, BRAND_GRAY3),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            el.append(t)
        add_section(_t("reports.ai_report.conclusion", lang), content.get("conclusion"))

    fname = f"{report_type}_{datetime.now().strftime('%Y%m%d')}.pdf"
    return _pdf_response(el, fname, brand, lang)


def _ai_report_excel(content: dict, report_type: str, lang: str = "es"):
    """Convierte el JSON de Claude a un Excel."""
    import openpyxl
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = _report_label(report_type, lang)[:31]

    hdr_fill = PatternFill("solid", fgColor="59008D")
    hdr_font = Font(color="FFFFFF", bold=True)
    thin = Side(style="thin", color="D1D5DB")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def write_section(title, text):
        if not text:
            return
        ws.append([title])
        ws.cell(row=ws.max_row, column=1).font = Font(bold=True, color="59008D", size=11)
        ws.append([str(text)])
        ws.cell(row=ws.max_row, column=1).alignment = Alignment(wrap_text=True)
        ws.append([""])

    def write_list(title, items):
        if not items:
            return
        ws.append([title])
        ws.cell(row=ws.max_row, column=1).font = Font(bold=True, color="59008D", size=11)
        for item in items:
            ws.append(["• " + str(item)])
        ws.append([""])

    ws.append([content.get("title", _report_label(report_type, lang))])
    ws.cell(row=1, column=1).font = Font(bold=True, size=14, color="59008D")
    ws.append([_t("reports.ai_report.org_date_label", lang, org=content.get("organization", ""), date=content.get("date", ""))])
    ws.append([""])

    write_section(_t("reports.ai_report.executive_summary", lang), content.get("executive_summary"))

    if report_type == "treatment_plan":
        write_section(_t("reports.ai_report.risk_appetite_analysis", lang), content.get("risk_appetite_analysis"))
        risks = content.get("risks", [])
        if risks:
            ws.append([_t("reports.ai_report.treatment_plans", lang)])
            ws.cell(row=ws.max_row, column=1).font = Font(bold=True, color="59008D", size=11)
            hdrs = [
                _t("reports.ai_report.header_code", lang), _t("reports.ai_report.header_priority", lang), _t("reports.ai_report.header_narrative", lang),
                _t("reports.ai_report.header_actions", lang), _t("reports.ai_report.header_metrics", lang), _t("reports.ai_report.header_effort", lang),
                _t("reports.ai_report.header_target_level", lang),
            ]
            ws.append(hdrs)
            for h_idx, _ in enumerate(hdrs, 1):
                c = ws.cell(row=ws.max_row, column=h_idx)
                c.fill = hdr_fill
                c.font = hdr_font
                c.border = border
            for r in risks:
                ws.append([
                    r.get("code", ""),
                    r.get("priority", ""),
                    r.get("treatment_narrative", ""),
                    "; ".join(r.get("recommended_actions") or []),
                    r.get("success_metrics", ""),
                    r.get("estimated_effort", ""),
                    r.get("target_residual_level", ""),
                ])
                for col_i in range(1, len(hdrs) + 1):
                    ws.cell(row=ws.max_row, column=col_i).border = border
        write_section(_t("reports.ai_report.implementation_roadmap_short", lang), content.get("implementation_roadmap"))
        write_section(_t("reports.ai_report.conclusion", lang), content.get("conclusion"))

    elif report_type == "executive_dashboard":
        write_list(_t("reports.ai_report.key_findings", lang), content.get("key_findings"))
        write_section(_t("reports.ai_report.risk_posture", lang), content.get("risk_posture_explanation"))
        write_section(_t("reports.ai_report.critical_risks", lang), content.get("top_risks_narrative"))
        write_section(_t("reports.ai_report.control_effectiveness", lang), content.get("control_effectiveness"))
        write_list(_t("reports.ai_report.critical_actions", lang), content.get("critical_actions"))
        write_section(_t("reports.ai_report.kpi_analysis", lang), content.get("kpi_commentary"))
        # Stats
        meta = content.get("_meta", {})
        stats = meta.get("stats", {})
        if stats:
            ws.append([_t("reports.ai_report.stats", lang)])
            ws.cell(row=ws.max_row, column=1).font = Font(bold=True, color="59008D", size=11)
            for k, v in stats.items():
                ws.append([k.replace("_", " ").capitalize(), v])

    elif report_type == "committee_minutes":
        write_list(_t("reports.ai_report.agenda", lang), content.get("agenda"))
        write_section(_t("reports.ai_report.risk_register_review_short", lang), content.get("risk_register_review"))
        accepted = content.get("accepted_risks", [])
        if accepted:
            ws.append([_t("reports.ai_report.accepted_risks_short", lang)])
            ws.cell(row=ws.max_row, column=1).font = Font(bold=True, color="59008D", size=11)
            ws.append([_t("reports.ai_report.header_code", lang), _t("reports.ai_report.header_justification", lang)])
            for ar in accepted:
                ws.append([ar.get("code", ""), ar.get("rationale", "")])
        write_list(_t("reports.ai_report.decisions_adopted", lang), content.get("decisions"))
        actions = content.get("action_items", [])
        if actions:
            ws.append([_t("reports.ai_report.agreed_actions", lang)])
            ws.cell(row=ws.max_row, column=1).font = Font(bold=True, color="59008D", size=11)
            ws.append([_t("reports.ai_report.header_action", lang), _t("reports.ai_report.header_responsible", lang), _t("reports.ai_report.header_deadline", lang)])
            for a in actions:
                ws.append([a.get("action", ""), a.get("responsible", ""), a.get("deadline", "")])
        write_section(_t("reports.ai_report.closing", lang), content.get("closing_note"))

    elif report_type == "followup_report":
        cl12 = content.get("cl12_assessment", {})
        if cl12:
            ws.append([_t("reports.ai_report.cl12_assessment", lang)])
            ws.cell(row=ws.max_row, column=1).font = Font(bold=True, color="59008D", size=11)
            for k, v in cl12.items():
                ws.append([k.replace("_", " ").capitalize(), str(v)])
        write_list(_t("reports.ai_report.strengths", lang), content.get("strengths"))
        write_list(_t("reports.ai_report.weaknesses", lang), content.get("weaknesses"))
        recs = content.get("recommendations", [])
        if recs:
            ws.append([_t("reports.ai_report.recommendations", lang)])
            ws.cell(row=ws.max_row, column=1).font = Font(bold=True, color="59008D", size=11)
            ws.append([_t("reports.ai_report.header_area", lang), _t("reports.ai_report.header_recommendation", lang), _t("reports.ai_report.header_priority", lang), _t("reports.ai_report.header_iso_ref_full", lang)])
            for rec in recs:
                ws.append([rec.get("area", ""), rec.get("recommendation", ""),
                           rec.get("priority", ""), rec.get("iso_reference", "")])
        write_section(_t("reports.ai_report.conclusion", lang), content.get("conclusion"))

    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 70
    for col in ["C", "D", "E", "F", "G"]:
        ws.column_dimensions[col].width = 20

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    fname = f"{report_type}_{datetime.now().strftime('%Y%m%d')}.xlsx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={fname}"},
    )
