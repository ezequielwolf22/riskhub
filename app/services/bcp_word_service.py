"""Exportacion del modulo BCP a Word (.docx) editable.

El cliente quiere entregables que pueda abrir y retocar en Word: informes de
continuidad, no solo una hoja de datos. Este servicio genera un documento
formal, con la paleta de marca (purple/orange), a partir de los mismos datos
que muestra la aplicacion. Todo lo del modulo es exportable: BIA (con el metodo
declarado de la organizacion), escenarios de indisponibilidad, estrategias de
recuperacion (DRP), planes, pruebas, proveedores criticos y dependencias.

Principio: el Word REFLEJA el estado guardado, no calcula nada nuevo. Las cifras
de impacto ponderado y banda salen del motor determinista (`bcm_scenario_engine`)
igual que en pantalla; aqui solo se maquetan.

python-docx ya es dependencia (se usa en la ingesta para leer .docx). El
documento se devuelve como bytes para servirlo desde el router.
"""
from __future__ import annotations

import io
import logging
from datetime import datetime, timezone
from typing import Callable, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from sqlalchemy.orm import Session

logger = logging.getLogger("riskhub.bcp_word")

# Paleta de marca (misma que el resto de informes RiskHub).
PURPLE = RGBColor(0x59, 0x00, 0x8D)
ORANGE = RGBColor(0xD6, 0x52, 0x00)
GREY = RGBColor(0x6B, 0x72, 0x80)
_PURPLE_HEX = "59008D"
_LIGHT_HEX = "F3EEF7"

# Todas las secciones exportables, en el orden en que aparecen en el informe.
ALL_SECTIONS = ("status", "bia", "scenarios", "strategies", "plans",
                "tests", "suppliers", "dependencies")

_CRIT_LABELS = {"critical": "Critico", "high": "Alto",
                "medium": "Medio", "low": "Bajo"}
_BAND_LABELS = {"none": "Sin impacto", "trivial": "Trivial",
                "relevant": "Relevante", "severe": "Severo",
                "critical": "Critico"}
_STATUS_LABELS = {"draft": "Borrador", "approved": "Aprobado",
                  "under_review": "En revision", "active": "Activo",
                  "retired": "Retirado"}
_IMPL_LABELS = {"planned": "Planificada", "in_progress": "En curso",
                "implemented": "Implantada", "tested": "Probada"}


# ── Helpers de maquetacion ────────────────────────────────────────────────────

def _shade(cell, hex_color: str) -> None:
    """Fondo de una celda (python-docx no lo expone directamente)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = PURPLE if level == 1 else ORANGE
    run.font.size = Pt(16 if level == 1 else 13)
    p.space_before = Pt(12)
    p.space_after = Pt(4)


def _muted(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = GREY


def _fmt(value, dash: str = "—") -> str:
    if value is None or value == "":
        return dash
    if isinstance(value, bool):
        return "Si" if value else "No"
    if isinstance(value, float):
        return ("%.2f" % value).rstrip("0").rstrip(".")
    if isinstance(value, (list, tuple)):
        return ", ".join(str(v) for v in value) if value else dash
    return str(value)


def _date(iso: Optional[str]) -> str:
    if not iso:
        return "—"
    return str(iso)[:10]


def _table(doc: Document, headers: list[str], rows: list[list],
           widths: Optional[list[float]] = None) -> None:
    """Tabla con cabecera de marca. `rows` puede ir vacia."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _shade(hdr[i], _PURPLE_HEX)
        para = hdr[i].paragraphs[0]
        run = para.add_run(str(h))
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if i >= len(cells):
                break
            para = cells[i].paragraphs[0]
            run = para.add_run(_fmt(val))
            run.font.size = Pt(9)
        if r_idx % 2 == 1:
            for c in cells:
                _shade(c, _LIGHT_HEX)
    doc.add_paragraph()


def _kv(doc: Document, pairs: list[tuple[str, object]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in pairs:
        cells = table.add_row().cells
        _shade(cells[0], _LIGHT_HEX)
        run = cells[0].paragraphs[0].add_run(str(label))
        run.bold = True
        run.font.size = Pt(9)
        cells[1].paragraphs[0].add_run(_fmt(value)).font.size = Pt(9)
    doc.add_paragraph()


def _empty(doc: Document, msg: str = "Sin registros.") -> None:
    p = doc.add_paragraph()
    run = p.add_run(msg)
    run.italic = True
    run.font.color.rgb = GREY


def _names_by_id(db: Session, org: int) -> dict:
    from app.models import User
    return {u.id: (u.full_name or u.email)
            for u in db.query(User).filter_by(organization_id=org).all()}


def _proc_names(db: Session, org: int) -> dict:
    from app.models import BusinessProcess
    return {p.id: p.name for p in db.query(BusinessProcess).filter_by(
        organization_id=org).all()}


def _proc_names_from_ids(ids, pmap: dict) -> str:
    if not ids:
        return "—"
    return ", ".join(pmap.get(i, "#%s" % i) for i in ids)


# ── Secciones ─────────────────────────────────────────────────────────────────

def _section_status(doc: Document, db: Session, org: int) -> None:
    from app.routers.bcp import bcp_dashboard  # reutiliza el mismo calculo
    _heading(doc, "Estado del SGCN", 1)
    _muted(doc, "Indicadores de continuidad de negocio (ISO 22301).")
    # Se invoca la logica del dashboard sin la capa HTTP.
    from app.models import BCPPlan, BCPTest, BusinessProcess
    from app.services.bcp_service import bia_completeness
    procs = db.query(BusinessProcess).filter_by(organization_id=org).all()
    tests = db.query(BCPTest).filter_by(organization_id=org).all()
    plans = db.query(BCPPlan).filter_by(organization_id=org).all()
    now = datetime.now(timezone.utc)
    critical = [p for p in procs if p.criticality == "critical"]
    overdue = [p for p in procs if not p.last_tested_at or
               (now - p.last_tested_at.replace(tzinfo=timezone.utc)).days > 365]
    tests_done = [t for t in tests if t.conducted_at]
    approved = [p for p in plans if p.status == "approved"]
    bia_pcts = [bia_completeness(None, p)["pct"] for p in procs] if procs else []
    bia_avg = int(sum(bia_pcts) / len(bia_pcts)) if bia_pcts else 0
    _kv(doc, [
        ("Procesos registrados", len(procs)),
        ("Procesos criticos", len(critical)),
        ("BIA completado (media)", "%d%%" % bia_avg),
        ("Procesos sin probar (>12m)", len(overdue)),
        ("Pruebas registradas", len(tests)),
        ("Pruebas realizadas", len(tests_done)),
        ("Planes aprobados", len(approved)),
    ])


def _section_bia(doc: Document, db: Session, org: int) -> None:
    from app.models import BusinessProcess
    from app.services.bcm_scenario_engine import get_criteria, scenario_matrix
    from app.services.bcp_service import bia_completeness
    _heading(doc, "Analisis de Impacto en el Negocio (BIA)", 1)

    # Metodo declarado por la organizacion: dimensiones, baremo RTO y bandas.
    criteria = get_criteria(db, org)
    _heading(doc, "Metodo de valoracion", 2)
    _muted(doc, "Dimensiones, baremo de RTO y bandas con los que esta "
                "organizacion valora el impacto. Editable en la aplicacion.")
    dims = criteria.get("dimensions") or []
    _table(doc, ["Dimension", "Clave"],
           [[d.get("label"), d.get("key")] for d in dims])
    rto = criteria.get("rto_scale") or []
    if rto:
        _heading(doc, "Baremo de RTO", 2)
        _table(doc, ["RTO", "Horas", "Factor"],
               [[r.get("label"), r.get("hours"), r.get("factor")] for r in rto])
    bands = criteria.get("bands") or []
    if bands:
        _heading(doc, "Bandas de impacto", 2)
        _table(doc, ["Banda", "Desde", "Hasta"],
               [[_BAND_LABELS.get(b.get("key"), b.get("label")),
                 b.get("min"), b.get("max")] for b in bands])
    comb = (criteria.get("combination") or "product")
    comb_txt = ("RTO + criterio = impacto total (suma)" if comb == "sum"
                else "impacto x factor de RTO (producto)" if comb == "product"
                else "formula propia")
    _muted(doc, "Combinacion impacto/RTO: %s." % comb_txt)

    # Procesos criticos y sus objetivos de recuperacion.
    _heading(doc, "Procesos y objetivos de recuperacion", 2)
    procs = db.query(BusinessProcess).filter_by(organization_id=org).order_by(
        BusinessProcess.criticality).all()
    if not procs:
        _empty(doc)
    else:
        rows = []
        for p in procs:
            bia = bia_completeness(None, p)
            rows.append([
                p.name, _CRIT_LABELS.get(p.criticality, p.criticality),
                _fmt(p.rto_hours), _fmt(p.rpo_hours), _fmt(p.mtpd_hours),
                _fmt(getattr(p, "weighted_impact", None)),
                _BAND_LABELS.get(getattr(p, "impact_band", None),
                                 getattr(p, "impact_band", None) or "—"),
                "%d%%" % bia["pct"],
            ])
        _table(doc, ["Proceso", "Criticidad", "RTO (h)", "RPO (h)",
                     "MTPD (h)", "Impacto", "Banda", "BIA"], rows)

    # Cobertura por escenario de indisponibilidad.
    matrix = scenario_matrix(db, org)
    if matrix.get("scenarios"):
        _heading(doc, "Cobertura de escenarios de indisponibilidad", 2)
        _muted(doc, "Cobertura valorada: %d%% (%d de %d aplicables)." % (
            matrix.get("coverage_pct", 0), matrix.get("assessed_total", 0),
            matrix.get("applicable_total", 0)))


def _section_scenarios(doc: Document, db: Session, org: int) -> None:
    from app.models import BCMLocation, BCMScenario, BCMScenarioAssessment
    _heading(doc, "Escenarios de indisponibilidad", 1)
    _muted(doc, "Valoracion por escenario y sede (impacto ponderado y banda "
                "del motor determinista, ISO 22301 cl. 8.2).")
    scen = {s.id: s for s in db.query(BCMScenario).filter_by(
        organization_id=org).all()}
    locs = {loc.id: loc.name for loc in db.query(BCMLocation).filter_by(
        organization_id=org).all()}
    assessments = db.query(BCMScenarioAssessment).filter_by(
        organization_id=org).all()
    if not assessments:
        _empty(doc, "No hay valoraciones de escenarios registradas.")
        return
    rows = []
    for a in assessments:
        s = scen.get(a.scenario_id)
        rows.append([
            s.code if s else "—",
            s.name if s else "—",
            locs.get(a.location_id, "Global"),
            _fmt(a.rto_label or a.rto_hours),
            _fmt(a.weighted_impact),
            _BAND_LABELS.get(a.impact_band, a.impact_band or "—"),
        ])
    _table(doc, ["Codigo", "Escenario", "Sede", "RTO", "Impacto", "Banda"], rows)


def _section_strategies(doc: Document, db: Session, org: int) -> None:
    from app.models import BCPStrategy
    _heading(doc, "Estrategias de recuperacion (DRP)", 1)
    _muted(doc, "Estrategias de recuperacion y planes de recuperacion ante "
                "desastres por proceso.")
    pmap = _proc_names(db, org)
    strats = db.query(BCPStrategy).filter_by(organization_id=org).all()
    if not strats:
        _empty(doc)
        return
    rows = [[
        s.name, s.strategy_type or "—",
        pmap.get(s.process_id, "Global"),
        _IMPL_LABELS.get(s.implementation_status, s.implementation_status or "—"),
        _fmt(s.estimated_cost), _date(s.target_date.isoformat() if s.target_date else None),
    ] for s in strats]
    _table(doc, ["Estrategia", "Tipo", "Proceso", "Estado", "Coste est.",
                 "Fecha objetivo"], rows)


def _section_plans(doc: Document, db: Session, org: int) -> None:
    from app.models import BCPPlan
    _heading(doc, "Planes de continuidad", 1)
    plans = db.query(BCPPlan).filter_by(organization_id=org).all()
    if not plans:
        _empty(doc)
        return
    rows = [[
        p.code or "—", p.name, p.plan_type or "—", p.version or "—",
        _STATUS_LABELS.get(p.status, p.status or "—"),
        _date(p.review_date.isoformat() if p.review_date else None),
    ] for p in plans]
    _table(doc, ["Codigo", "Plan", "Tipo", "Version", "Estado", "Revision"], rows)


def _section_tests(doc: Document, db: Session, org: int) -> None:
    from app.models import BCPTest
    _heading(doc, "Pruebas de continuidad", 1)
    _muted(doc, "Ejercicios y pruebas de continuidad (ISO 22301 cl. 8.5).")
    tests = db.query(BCPTest).filter_by(organization_id=org).order_by(
        BCPTest.conducted_at.desc().nullslast()).all()
    if not tests:
        _empty(doc)
        return
    rows = [[
        t.code or "—", t.test_type or "—",
        _date(t.conducted_at.isoformat() if t.conducted_at else None),
        t.result or "—",
        _fmt(t.rto_achieved_hours),
    ] for t in tests]
    _table(doc, ["Codigo", "Tipo", "Fecha", "Resultado", "RTO logrado (h)"], rows)


def _section_suppliers(doc: Document, db: Session, org: int) -> None:
    from app.models import BCPSupplierLink, Supplier
    _heading(doc, "Proveedores criticos", 1)
    _muted(doc, "Proveedores relevantes para la continuidad y su contingencia.")
    links = db.query(BCPSupplierLink).filter_by(organization_id=org).all()
    if not links:
        _empty(doc)
        return
    rows = []
    for s in links:
        sup = db.get(Supplier, s.supplier_id) if s.supplier_id else None
        alt = db.get(Supplier, s.alternative_supplier_id) if s.alternative_supplier_id else None
        rows.append([
            sup.name if sup else "—",
            _CRIT_LABELS.get(s.criticality, s.criticality or "—"),
            _fmt(s.rto_impact_hours),
            "Si" if s.has_contingency_plan else "No",
            alt.name if alt else "—",
        ])
    _table(doc, ["Proveedor", "Criticidad", "Impacto RTO (h)",
                 "Plan contingencia", "Proveedor alternativo"], rows)


def _section_dependencies(doc: Document, db: Session, org: int) -> None:
    from app.models import BCPDependency
    _heading(doc, "Dependencias", 1)
    _muted(doc, "Recursos y dependencias que necesita cada proceso para "
                "recuperarse.")
    pmap = _proc_names(db, org)
    deps = db.query(BCPDependency).filter_by(organization_id=org).all()
    if not deps:
        _empty(doc)
        return
    rows = [[
        pmap.get(d.process_id, "#%s" % d.process_id),
        d.dependency_type or "—", d.name or "—",
        "Si" if d.is_critical else "No",
        _fmt(d.rto_hours), d.alternative or "—",
    ] for d in deps]
    _table(doc, ["Proceso", "Tipo", "Recurso", "Critica", "RTO (h)",
                 "Alternativa"], rows)


_BUILDERS: dict[str, Callable] = {
    "status": _section_status,
    "bia": _section_bia,
    "scenarios": _section_scenarios,
    "strategies": _section_strategies,
    "plans": _section_plans,
    "tests": _section_tests,
    "suppliers": _section_suppliers,
    "dependencies": _section_dependencies,
}

_SECTION_TITLES = {
    "status": "Estado", "bia": "BIA", "scenarios": "Escenarios",
    "strategies": "Estrategias (DRP)", "plans": "Planes", "tests": "Pruebas",
    "suppliers": "Proveedores criticos", "dependencies": "Dependencias",
}


# ── Portada y entrada publica ─────────────────────────────────────────────────

def _cover(doc: Document, org_name: str, sections: list[str]) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Plan de Continuidad de Negocio")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = PURPLE
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(org_name)
    r2.font.size = Pt(16)
    r2.font.color.rgb = ORANGE
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Generado el %s" % datetime.now().strftime("%d/%m/%Y"))
    r3.font.size = Pt(11)
    r3.font.color.rgb = GREY
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    labels = ", ".join(_SECTION_TITLES.get(s, s) for s in sections)
    r4 = p4.add_run("Contenido: %s" % labels)
    r4.font.size = Pt(10)
    r4.font.color.rgb = GREY
    doc.add_page_break()


def normalize_sections(raw: Optional[str]) -> list[str]:
    """Convierte el parametro `sections` en una lista valida y ordenada."""
    if not raw or raw.strip().lower() in ("all", "todo", "todos", "*"):
        return list(ALL_SECTIONS)
    wanted = {s.strip().lower() for s in raw.split(",") if s.strip()}
    ordered = [s for s in ALL_SECTIONS if s in wanted]
    return ordered or list(ALL_SECTIONS)


def generate_bcp_word(db: Session, org: int, org_name: str = "",
                      sections: Optional[list[str]] = None) -> bytes:
    """Genera el informe BCP en Word y devuelve los bytes del .docx."""
    sects = sections or list(ALL_SECTIONS)
    doc = Document()
    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(10)

    _cover(doc, org_name or "Organizacion", sects)
    for key in sects:
        builder = _BUILDERS.get(key)
        if not builder:
            continue
        try:
            builder(doc, db, org)
        except Exception:
            logger.exception("bcp_word: fallo generando la seccion %s", key)
            _empty(doc, "No se pudo generar esta seccion.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
