"""Lectura de documentos preservando su estructura.

Los extractores de `document_service` aplanan el documento a una cadena de
texto. Para un informe eso vale; para comprender un BIA no. El libro de un
cliente real trae once hojas con estructuras distintas: escenarios en bloques
de siete filas, aplicaciones criticas fila a fila, y los baremos como matriz.
Aplanar eso destruye justo la informacion que dice como descomponerlo.

Aqui el documento se devuelve como una lista de bloques con su referencia de
origen (`sheet:Escenarios-Ranking`, `h1:4`, `table:2`). Esa referencia viaja
despues hasta el mapa de volcado y hasta el registro de conflictos, para que
cualquier dato importado se pueda rastrear hasta la celda de la que salio.

Este modulo no importa `document_service` ni `isms_analysis_service` a
proposito: son piezas de otro pipeline con su propio ciclo de vida.
"""
from __future__ import annotations

import hashlib
import io
import logging
import re
from typing import Any, Optional

logger = logging.getLogger("riskhub.ingest.reader")

# Topes de seguridad: un libro corrupto o generado por maquina no debe agotar
# la memoria del worker. Se recorta y se deja constancia en `truncated`.
MAX_SHEET_ROWS = 400
MAX_SHEET_COLS = 40
MAX_BLOCKS = 400
MAX_CELL_CHARS = 2000
MAX_TEXT_CHARS = 20000

_SUPPORTED = (".xlsx", ".xlsm", ".xls", ".docx", ".pdf", ".txt", ".csv", ".md")


class UnsupportedDocument(ValueError):
    """El formato no se puede leer con las herramientas disponibles."""


def sha256_of(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def is_supported(filename: str) -> bool:
    return str(filename or "").lower().endswith(_SUPPORTED)


def read_document(data: bytes, filename: str) -> dict:
    """Devuelve la estructura del documento como bloques con referencia.

    {filename, sha256, format, blocks: [...], truncated: bool, stats: {...}}
    """
    name = str(filename or "").lower()
    if name.endswith((".xlsx", ".xlsm", ".xls")):
        fmt, blocks, truncated = "xlsx", *_read_xlsx(data)
    elif name.endswith(".docx"):
        fmt, blocks, truncated = "docx", *_read_docx(data)
    elif name.endswith(".pdf"):
        fmt, blocks, truncated = "pdf", *_read_pdf(data)
    elif name.endswith(".csv"):
        fmt, blocks, truncated = "csv", *_read_csv(data)
    elif name.endswith((".txt", ".md")):
        fmt, blocks, truncated = "text", *_read_plain(data)
    else:
        raise UnsupportedDocument(f"Formato no soportado: {filename}")

    return {
        "filename": filename,
        "sha256": sha256_of(data),
        "format": fmt,
        "blocks": blocks[:MAX_BLOCKS],
        "truncated": truncated or len(blocks) > MAX_BLOCKS,
        "stats": {
            "blocks": min(len(blocks), MAX_BLOCKS),
            "sheets": sum(1 for b in blocks if b["type"] == "sheet"),
            "tables": sum(1 for b in blocks if b["type"] == "table"),
            "sections": sum(1 for b in blocks if b["type"] == "section"),
        },
    }


# ── XLSX ─────────────────────────────────────────────────────────────────────

def _cell(value: Any) -> str:
    """Normaliza una celda a texto. Los errores de formula no son datos."""
    if value is None:
        return ""
    text = str(value).strip()
    if text in ("#REF!", "#N/A", "#VALUE!", "#DIV/0!", "#NAME?", "#NULL!", "#NUM!"):
        return ""
    return text[:MAX_CELL_CHARS]


def _read_xlsx(data: bytes) -> tuple[list[dict], bool]:
    import openpyxl

    # read_only=False para poder consultar celdas combinadas: en las plantillas
    # de BIA la combinacion marca el alcance de un bloque, no es decoracion.
    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
    blocks: list[dict] = []
    truncated = False

    for ws in wb.worksheets:
        max_row = min(ws.max_row or 0, MAX_SHEET_ROWS)
        max_col = min(ws.max_column or 0, MAX_SHEET_COLS)
        if (ws.max_row or 0) > MAX_SHEET_ROWS or (ws.max_column or 0) > MAX_SHEET_COLS:
            truncated = True
        if not max_row or not max_col:
            continue

        grid: list[list[str]] = []
        for row in ws.iter_rows(min_row=1, max_row=max_row, max_col=max_col,
                                values_only=True):
            grid.append([_cell(c) for c in row])

        # Se descartan filas y columnas totalmente vacias, pero se conserva el
        # indice real de cada fila para que las referencias sigan apuntando a
        # la celda de origen del libro.
        row_index = [i for i, r in enumerate(grid, 1) if any(c for c in r)]
        if not row_index:
            continue
        used_cols = sorted({
            j for i in row_index for j, c in enumerate(grid[i - 1]) if c
        })
        compact = [[grid[i - 1][j] if j < len(grid[i - 1]) else "" for j in used_cols]
                   for i in row_index]

        merged = []
        try:
            for rng in list(ws.merged_cells.ranges)[:200]:
                merged.append(str(rng))
        except Exception:
            pass

        blocks.append({
            "type": "sheet",
            "ref": f"sheet:{ws.title}",
            "name": ws.title,
            "columns": [_col_letter(j + 1) for j in used_cols],
            "row_numbers": row_index,
            "grid": compact,
            "merged_ranges": merged,
        })

    wb.close()
    return blocks, truncated


def _col_letter(idx: int) -> str:
    letters = ""
    while idx > 0:
        idx, rem = divmod(idx - 1, 26)
        letters = chr(65 + rem) + letters
    return letters


# ── DOCX ─────────────────────────────────────────────────────────────────────

def _read_docx(data: bytes) -> tuple[list[dict], bool]:
    import docx
    from docx.document import Document as _Doc
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table as _Table
    from docx.text.paragraph import Paragraph as _Para

    doc = docx.Document(io.BytesIO(data))
    blocks: list[dict] = []
    truncated = False

    def _iter_body(parent):
        """Parrafos y tablas en el orden real del documento.

        python-docx expone `paragraphs` y `tables` por separado, lo que pierde
        el orden: una tabla dejaria de saberse a que epigrafe pertenece.
        """
        body = parent.element.body if isinstance(parent, _Doc) else parent._tc
        for child in body.iterchildren():
            if isinstance(child, CT_P):
                yield _Para(child, parent)
            elif isinstance(child, CT_Tbl):
                yield _Table(child, parent)

    section_idx = 0
    table_idx = 0
    current: Optional[dict] = None
    para_idx = 0

    for item in _iter_body(doc):
        if isinstance(item, _Para):
            para_idx += 1
            text = item.text.strip()
            if not text:
                continue
            style = (item.style.name if item.style is not None else "") or ""
            level = _heading_level(style)
            if level:
                section_idx += 1
                current = {
                    "type": "section", "ref": f"h{level}:{section_idx}",
                    "title": text[:300], "level": level, "text": "",
                }
                blocks.append(current)
            elif current is not None:
                if len(current["text"]) < MAX_TEXT_CHARS:
                    current["text"] += (text + "\n")
                else:
                    truncated = True
            else:
                blocks.append({"type": "text", "ref": f"p:{para_idx}",
                               "text": text[:MAX_TEXT_CHARS]})
        else:
            table_idx += 1
            rows = []
            for r in item.rows[:MAX_SHEET_ROWS]:
                rows.append([c.text.strip()[:MAX_CELL_CHARS]
                             for c in r.cells[:MAX_SHEET_COLS]])
            if not rows:
                continue
            blocks.append({
                "type": "table", "ref": f"table:{table_idx}",
                "section": current["title"] if current else None,
                "header": rows[0],
                "rows": rows[1:],
            })

    return blocks, truncated


def _heading_level(style_name: str) -> int:
    """Nivel de encabezado. Contempla los nombres localizados de Word."""
    s = style_name.strip().lower()
    m = re.match(r"^(heading|titulo|título|t[ií]tulo)\s*(\d+)$", s)
    if m:
        return min(6, int(m.group(2)))
    if s in ("title", "titulo", "título"):
        return 1
    return 0


# ── PDF ──────────────────────────────────────────────────────────────────────

def _read_pdf(data: bytes) -> tuple[list[dict], bool]:
    import pypdf

    reader = pypdf.PdfReader(io.BytesIO(data))
    blocks: list[dict] = []
    for i, page in enumerate(reader.pages[:200], 1):
        try:
            text = (page.extract_text() or "").strip()
        except Exception:
            text = ""
        if text:
            blocks.append({"type": "page", "ref": f"page:{i}", "page": i,
                           "text": text[:MAX_TEXT_CHARS]})
    # Un PDF sin texto extraible es casi siempre un escaneo: no es un error,
    # es un caso para Vision. Se marca para que el llamador lo derive.
    if not blocks:
        blocks.append({"type": "scanned", "ref": "page:1",
                       "text": "", "needs_vision": True})
    return blocks, False


# ── CSV y texto plano ────────────────────────────────────────────────────────

def _decode(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _read_csv(data: bytes) -> tuple[list[dict], bool]:
    import csv

    text = _decode(data)
    try:
        dialect = csv.Sniffer().sniff(text[:4096], delimiters=",;\t|")
    except Exception:
        dialect = csv.excel
    rows = [
        [str(c).strip()[:MAX_CELL_CHARS] for c in row]
        for row in csv.reader(text.splitlines(), dialect)
    ]
    rows = [r for r in rows if any(r)]
    if not rows:
        return [], False
    truncated = len(rows) > MAX_SHEET_ROWS
    rows = rows[:MAX_SHEET_ROWS]
    return [{"type": "table", "ref": "table:1", "section": None,
             "header": rows[0], "rows": rows[1:]}], truncated


def _read_plain(data: bytes) -> tuple[list[dict], bool]:
    text = _decode(data)
    truncated = len(text) > MAX_TEXT_CHARS
    return [{"type": "text", "ref": "p:1", "text": text[:MAX_TEXT_CHARS]}], truncated


# ── Serializacion para el prompt ─────────────────────────────────────────────

def render_for_llm(document: dict, max_chars: int = 60000) -> str:
    """Convierte la estructura a texto legible por el modelo, con referencias.

    Se conservan las referencias de bloque (`sheet:X`, `h1:3`) y los numeros de
    fila reales: son lo que permite al agente decir "las filas 8, 15 y 22 de la
    hoja Escenarios-Ranking son tres unidades" en vez de inventarse el reparto.
    """
    parts: list[str] = [f"### DOCUMENTO: {document.get('filename')}"]
    budget = max_chars

    for block in document.get("blocks", []):
        if budget <= 0:
            parts.append("\n[...documento recortado por longitud...]")
            break
        btype = block["type"]
        if btype == "sheet":
            head = [f"\n## HOJA \"{block['name']}\"  [ref={block['ref']}]"]
            head.append("columnas: " + ", ".join(block.get("columns") or []))
            if block.get("merged_ranges"):
                head.append("celdas combinadas: "
                            + ", ".join(block["merged_ranges"][:30]))
            lines = head
            for row_no, row in zip(block.get("row_numbers") or [],
                                   block.get("grid") or []):
                if not any(row):
                    continue
                lines.append(f"f{row_no}: " + " | ".join(row))
            chunk = "\n".join(lines)
        elif btype == "section":
            chunk = (f"\n## {'#' * block['level']} {block['title']}  "
                     f"[ref={block['ref']}]\n{block.get('text', '')}")
        elif btype == "table":
            lines = [f"\n## TABLA  [ref={block['ref']}]"]
            if block.get("section"):
                lines.append(f"(dentro de: {block['section']})")
            lines.append("cabecera: " + " | ".join(block.get("header") or []))
            for i, row in enumerate(block.get("rows") or [], 1):
                lines.append(f"f{i}: " + " | ".join(row))
            chunk = "\n".join(lines)
        elif btype == "page":
            chunk = f"\n## PAGINA {block['page']}  [ref={block['ref']}]\n{block['text']}"
        elif btype == "scanned":
            chunk = "\n[documento escaneado sin texto extraible]"
        else:
            chunk = f"\n{block.get('text', '')}"

        chunk = chunk[:budget]
        budget -= len(chunk)
        parts.append(chunk)

    return "\n".join(parts)
