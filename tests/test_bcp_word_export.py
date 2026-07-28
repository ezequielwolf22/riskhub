"""Export del modulo BCP a Word (.docx).

Verifica que el informe se genera como un .docx valido, que refleja el metodo
declarado de la organizacion (BIA) y que se puede pedir por secciones.
"""
import io

from docx import Document


def _docx_text(content: bytes) -> str:
    assert content[:2] == b"PK", "un .docx es un zip (empieza por PK)"
    doc = Document(io.BytesIO(content))
    return "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )


def test_export_word_completo_es_un_docx_valido(client, auth_headers):
    client.post("/api/bcp/processes", headers=auth_headers, json={
        "name": "Facturacion electronica", "criticality": "critical",
        "rto_hours": 4, "rpo_hours": 1, "mtpd_hours": 8,
    })
    resp = client.get("/api/bcp/export/word", headers=auth_headers)
    assert resp.status_code == 200, resp.text
    assert "wordprocessingml" in resp.headers["content-type"]
    assert "Informe_BCP.docx" in resp.headers.get("content-disposition", "")
    text = _docx_text(resp.content)
    assert "Plan de Continuidad de Negocio" in text
    assert "Analisis de Impacto en el Negocio" in text
    assert "Facturacion electronica" in text


def test_export_word_por_seccion(client, auth_headers):
    resp = client.get("/api/bcp/export/word?sections=bia", headers=auth_headers)
    assert resp.status_code == 200, resp.text
    assert "BCP_bia.docx" in resp.headers.get("content-disposition", "")
    text = _docx_text(resp.content)
    # La seccion BIA muestra el metodo declarado (dimensiones del cliente).
    assert "Metodo de valoracion" in text
    # No debe arrastrar otras secciones cuando se pide solo una.
    assert "Proveedores criticos" not in text


def test_export_word_refleja_las_dimensiones_del_metodo(client, auth_headers):
    # Cambia el metodo de la organizacion y comprueba que el Word lo refleja.
    client.put("/api/bcp/bia-criteria", headers=auth_headers, json={
        "dimensions": [
            {"key": "operational", "label": "Operativo"},
            {"key": "safety", "label": "Seguridad de las personas"},
        ],
        "horizons": ["0h", ">4h"],
        "levels": [
            {"value": 1, "score": 0, "label": "Nulo"},
            {"value": 2, "score": 2, "label": "Alto"},
        ],
        "bands": [
            {"key": "none", "label": "Sin impacto", "min": 0.0, "max": 0.5},
            {"key": "critical", "label": "Critico", "min": 0.5, "max": 4.0},
        ],
        "aggregation": "max",
    })
    resp = client.get("/api/bcp/export/word?sections=bia", headers=auth_headers)
    assert resp.status_code == 200, resp.text
    text = _docx_text(resp.content)
    assert "Seguridad de las personas" in text
